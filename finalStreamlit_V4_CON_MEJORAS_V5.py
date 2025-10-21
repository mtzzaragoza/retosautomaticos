import streamlit as st
import pandas as pd
from docx import Document
import io
import re
import random
import streamlit.components.v1 as components
from bs4 import BeautifulSoup
from datetime import datetime
import json
import os
from difflib import SequenceMatcher

# Intentar importar librerías de PDF
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Configuración de la página
st.set_page_config(page_title="Sistema de Retroalimentación", layout="wide")

# Sidebar para navegación
st.sidebar.title("📚 Sistema de Retroalimentación")
menu_option = st.sidebar.selectbox(
    "Selecciona una opción:",
    ["R3MD - Conjuntos", "R4MD - Proposiciones Lógicas", "R7MD - Mensajes Predefinidos"]
)

# ==================== FUNCIONES COMPARTIDAS ====================

def copy_to_clipboard_js(text):
    """Genera JavaScript para copiar texto al portapapeles"""
    # Escapar backticks ANTES del f-string para evitar error de sintaxis
    text_escaped = text.replace('`', '\\`')
    js_code = f"""
    <script>
        navigator.clipboard.writeText(`{text_escaped}`).then(function() {{
            console.log('Texto copiado al portapapeles');
        }});
    </script>
    """
    return js_code

# ==================== R3MD - CONJUNTOS (VERSIÓN DEFINITIVA FUSIONADA) ====================

def extraer_texto_pdf(pdf_file):
    """Extrae texto de un archivo PDF usando pdfplumber"""
    if not PDF_AVAILABLE:
        raise Exception("Las librerías de PDF no están instaladas. Instala: pip install PyPDF2 pdfplumber")
    
    try:
        texto_completo = ""
        with pdfplumber.open(pdf_file) as pdf:
            for pagina in pdf.pages:
                texto_pagina = pagina.extract_text()
                if texto_pagina:
                    texto_completo += texto_pagina + "\n"
        return texto_completo.strip()
    except Exception as e:
        try:
            pdf_file.seek(0)
            texto_completo = ""
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for pagina in pdf_reader.pages:
                texto_completo += pagina.extract_text() + "\n"
            return texto_completo.strip()
        except Exception as e2:
            raise Exception(f"Error con pdfplumber: {str(e)} | Error con PyPDF2: {str(e2)}")

def extraer_texto_docx_completo(docx_file):
    """Extrae texto de un archivo DOCX incluyendo párrafos y tablas en orden"""
    doc = Document(docx_file)
    
    # Crear un diccionario para mantener el orden de elementos
    elementos = []
    
    # Extraer párrafos con su índice
    for para in doc.paragraphs:
        if para.text.strip():
            elementos.append(('parrafo', para.text))
    
    # Extraer tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    elementos.append(('tabla', cell.text))
    
    # Unir todo el texto
    texto_completo = "\n".join([elem[1] for elem in elementos])
    return texto_completo, doc  # También retornar el objeto doc

# ========== FUNCIONES PARA WORD (de app_mejorado.py) ==========

def extraer_numeros_de_texto(texto):
    """
    Extrae números de un texto, soportando:
    - Separados por comas: 1,2,3
    - Separados por espacios: 1 2 3
    - Combinaciones: 1, 2, 3 o 1,2, 3
    Retorna un conjunto normalizado de números como strings
    """
    # Extraer todos los números
    numeros = re.findall(r'\d+', texto)
    
    # Normalizar (convertir a int y luego a string para eliminar ceros a la izquierda)
    numeros_normalizados = set(str(int(num)) for num in numeros if num)
    
    return numeros_normalizados

def extraer_respuestas_desde_doc(doc):
    """
    MÉTODO PARA WORD: Extrae TODAS las respuestas de un documento Word, 
    saltando el ejemplo inicial y los cálculos intermedios. 
    VERSIÓN V5 MEJORADA: 
    - Soporta combinaciones mixtas de delimitadores
    - Filtra líneas con "R=" sin contexto de operación
    - Mejora la detección de respuestas válidas
    """
    respuestas = []
    primera_respuesta = True  # Flag para saltar el ejemplo
    
    # Conjuntos base para filtrar SOLO en contextos de definición
    CONJUNTOS_BASE = [
        {'1','2','3','4','5','6','7','8','9','10','11','12','13','14'},  # U
        {'2','4','6','8','10','12','14'},  # A
        {'1','2','3','5','8','13'},  # B
        {'1','2','4','6','7','10','11','13'},  # C
    ]
    
    for table_idx, table in enumerate(doc.tables):
        # Saltar las primeras 3 tablas (encabezado, instrucciones, conjuntos base)
        if table_idx < 3:
            continue
        
        for row in table.rows:
            for cell in row.cells:
                texto_celda = cell.text.strip()
                
                # FILTRO PRIORITARIO: Ignorar líneas que solo contienen "R=" sin contexto de operación
                # Verificar si la línea tiene "R=" pero NO tiene símbolos de conjuntos (A, B, C)
                # y NO viene después de "Resultado de la operación:"
                if re.match(r'^\s*R\s*=', texto_celda, re.IGNORECASE):
                    # Solo aceptar si tiene contexto de operación (contiene A, B, C en operaciones)
                    tiene_operacion_conjunto = bool(re.search(r'[ABC]\s*[–\-∩∪´\']', texto_celda))
                    # O si tiene "Resultado de la operación:" en el mismo texto
                    tiene_resultado_en_texto = "Resultado de la operación:" in texto_celda or "resultado de la operacion" in texto_celda.lower()
                    
                    if not tiene_operacion_conjunto and not tiene_resultado_en_texto:
                        # Esta es una línea "R=" suelta sin contexto, IGNORARLA COMPLETAMENTE
                        continue
                
                # ESTRATEGIA 1: Buscar "Resultado de la operación:"
                tiene_resultado = "Resultado de la operación:" in texto_celda or "Resultado de la operacion:" in texto_celda.lower()
                
                # ESTRATEGIA 2: Buscar líneas con operaciones de conjuntos
                # Contiene = y algún tipo de delimitador o números con comas
                # MEJORADO: Verificar que contenga símbolos de conjuntos (A, B, C) en la operación
                tiene_operacion = ('=' in texto_celda and 
                                 ('{' in texto_celda or '[' in texto_celda or '(' in texto_celda or 
                                  (re.search(r'\d+,\s*\d+', texto_celda) is not None)) and 
                                 'Sean los conjuntos' not in texto_celda and
                                 len(texto_celda) < 300)  # Líneas cortas son más probables de ser respuestas
                
                # ESTRATEGIA 3: Buscar líneas que SOLO contienen un conjunto válido
                # Ejemplo: "[1, 2, 3]" o "{1, 3, 5}" o "(1, 3, 5, 7, 9, 11, 13)"
                solo_conjunto = (re.match(r'^\s*[\[\{\(].*[\]\}\)]\s*$', texto_celda) is not None and
                                len(texto_celda) < 100)
                
                if tiene_resultado or tiene_operacion or solo_conjunto:
                    # CASO ESPECIAL: Si tiene "Resultado de la operación:" seguido de "=" pero sin conjunto visible
                    # Esto indica que el estudiante no respondió este inciso
                    if tiene_resultado and '=' in texto_celda:
                        parte_despues_igual = texto_celda.split('=')[-1].strip()
                        # Si después del = no hay números ni delimitadores de conjuntos
                        if not re.search(r'[\d\{\[\(]', parte_despues_igual):
                            # Saltar el ejemplo si es el primero
                            if primera_respuesta:
                                primera_respuesta = False
                                continue
                            # Agregar conjunto vacío para mantener el orden
                            respuestas.append(set())
                            continue
                    
                    # Si hay múltiples "=", tomar solo lo que está después del último
                    if texto_celda.count('=') > 1:
                        texto_a_procesar = texto_celda.split('=')[-1].strip()
                    else:
                        texto_a_procesar = texto_celda
                    
                    # Buscar TODOS los conjuntos en el texto
                    conjuntos_encontrados = []
                    
                    # Determinar si debemos filtrar conjuntos base
                    # NO filtrar si viene después de "Resultado de la operación:" o tiene indicadores de resultado
                    es_respuesta = tiene_resultado or tiene_operacion or '=' in texto_celda
                    
                    # ============ NUEVA SECCIÓN: COMBINACIONES MIXTAS DE DELIMITADORES ============
                    
                    # Buscar combinaciones mixtas de delimitadores
                    # Formato: abre{cierra donde abre y cierra pueden ser diferentes
                    
                    # { ) - abre con llave, cierra con paréntesis
                    for match in re.finditer(r'\{([^\}\)]*)\)', texto_a_procesar):
                        numeros = re.findall(r'\d+', match.group(1))
                        if numeros and len(numeros) >= 2:
                            conjunto = set(str(int(num)) for num in numeros)
                            if es_respuesta or conjunto not in CONJUNTOS_BASE:
                                conjuntos_encontrados.append(conjunto)
                    
                    # { ] - abre con llave, cierra con corchete
                    for match in re.finditer(r'\{([^\}\]]*)\]', texto_a_procesar):
                        numeros = re.findall(r'\d+', match.group(1))
                        if numeros and len(numeros) >= 2:
                            conjunto = set(str(int(num)) for num in numeros)
                            if es_respuesta or conjunto not in CONJUNTOS_BASE:
                                conjuntos_encontrados.append(conjunto)
                    
                    # [ } - abre con corchete, cierra con llave
                    for match in re.finditer(r'\[([^\]\}]*)\}', texto_a_procesar):
                        numeros = re.findall(r'\d+', match.group(1))
                        if numeros and len(numeros) >= 2:
                            conjunto = set(str(int(num)) for num in numeros)
                            if es_respuesta or conjunto not in CONJUNTOS_BASE:
                                conjuntos_encontrados.append(conjunto)
                    
                    # [ ) - abre con corchete, cierra con paréntesis
                    for match in re.finditer(r'\[([^\]\)]*)\)', texto_a_procesar):
                        numeros = re.findall(r'\d+', match.group(1))
                        if numeros and len(numeros) >= 2:
                            conjunto = set(str(int(num)) for num in numeros)
                            if es_respuesta or conjunto not in CONJUNTOS_BASE:
                                conjuntos_encontrados.append(conjunto)
                    
                    # ( } - abre con paréntesis, cierra con llave
                    for match in re.finditer(r'\(([^\)\}]*)\}', texto_a_procesar):
                        numeros = re.findall(r'\d+', match.group(1))
                        if numeros and len(numeros) >= 2:
                            conjunto = set(str(int(num)) for num in numeros)
                            if es_respuesta or conjunto not in CONJUNTOS_BASE:
                                conjuntos_encontrados.append(conjunto)
                    
                    # ( ] - abre con paréntesis, cierra con corchete
                    for match in re.finditer(r'\(([^\)\]]*)\]', texto_a_procesar):
                        numeros = re.findall(r'\d+', match.group(1))
                        if numeros and len(numeros) >= 2:
                            conjunto = set(str(int(num)) for num in numeros)
                            if es_respuesta or conjunto not in CONJUNTOS_BASE:
                                conjuntos_encontrados.append(conjunto)
                    
                    # ============ COMBINACIONES ESTÁNDAR (MISMOS DELIMITADORES) ============
                    
                    # Corchetes []
                    for match in re.finditer(r'\[([^\]]+)\]', texto_a_procesar):
                        numeros = re.findall(r'\d+', match.group(1))
                        if numeros and len(numeros) >= 2:
                            conjunto = set(str(int(num)) for num in numeros)
                            # SI es respuesta, NO filtrar por conjuntos base
                            if es_respuesta or conjunto not in CONJUNTOS_BASE:
                                conjuntos_encontrados.append(conjunto)
                    
                    # Llaves {}
                    for match in re.finditer(r'\{([^}]+)\}', texto_a_procesar):
                        numeros = re.findall(r'\d+', match.group(1))
                        if numeros and len(numeros) >= 2:
                            conjunto = set(str(int(num)) for num in numeros)
                            # SI es respuesta, NO filtrar por conjuntos base
                            if es_respuesta or conjunto not in CONJUNTOS_BASE:
                                conjuntos_encontrados.append(conjunto)
                    
                    # Paréntesis ()
                    for match in re.finditer(r'\(([^)]+)\)', texto_a_procesar):
                        contenido = match.group(1)
                        if not re.search(r':\d{2}|Real|Máx', contenido):
                            numeros = re.findall(r'\d+', contenido)
                            if numeros and len(numeros) >= 2:
                                conjunto = set(str(int(num)) for num in numeros)
                                # SI es respuesta, NO filtrar por conjuntos base
                                if es_respuesta or conjunto not in CONJUNTOS_BASE:
                                    conjuntos_encontrados.append(conjunto)
                    
                    # Si no se encontraron con delimitadores, buscar números sueltos después del último "="
                    if not conjuntos_encontrados and '=' in texto_a_procesar:
                        parte_numeros = texto_a_procesar.split('=')[-1].strip()
                        numeros = re.findall(r'\d+', parte_numeros)
                        if numeros and len(numeros) >= 2:
                            conjunto = set(str(int(num)) for num in numeros)
                            # NO filtrar aquí porque viene después de "=" (es respuesta)
                            conjuntos_encontrados.append(conjunto)
                    
                    # Tomar el ÚLTIMO conjunto encontrado (el resultado final)
                    if conjuntos_encontrados:
                        # Saltar el primer resultado (ejemplo)
                        if primera_respuesta:
                            primera_respuesta = False
                            continue
                        
                        # Agregar el último conjunto (resultado final)
                        respuestas.append(conjuntos_encontrados[-1])
    
    return respuestas

# ========== FUNCIONES PARA PDF (de app_ULTRA_MAXIMA.py) ==========

def extraer_todos_los_numeros(texto):
    """Extrae TODOS los números de un texto, sin importar el formato"""
    numeros = re.findall(r'\d+', texto)
    if numeros:
        return set(str(int(num)) for num in numeros)
    return set()

def extraer_conjunto_agresivo(texto):
    """
    VERSIÓN V5 MEJORADA: Extrae conjuntos de CUALQUIER formato.
    Detecta: [], {}, (), combinaciones mixtas, números con espacios, y toma el ÚLTIMO conjunto válido.
    NUEVO V5: Soporta combinaciones mixtas de delimitadores.
    """
    todos_conjuntos = []
    
    # PASO 1: Si hay múltiples "=", dividir y tomar solo lo que está después del último
    if texto.count('=') > 1:
        partes = texto.split('=')
        # Tomar la última parte (después del último "=")
        texto = partes[-1].strip()
    
    # PASO 2: Buscar en TODOS los formatos posibles y guardar todos los conjuntos
    
    # ============ COMBINACIONES MIXTAS DE DELIMITADORES ============
    
    # { ) - abre con llave, cierra con paréntesis
    matches = re.finditer(r'\{([^\}\)]*)\)', texto)
    for match in matches:
        conjunto = extraer_todos_los_numeros(match.group(1))
        if conjunto and len(conjunto) >= 2:
            if conjunto not in [
                {'1','2','3','4','5','6','7','8','9','10','11','12','13','14'},
                {'2','4','6','8','10','12','14'},
                {'1','2','3','5','8','13'},
                {'1','2','4','6','7','10','11','13'},
                {'4','6','7','9','10','11','12','14'}
            ]:
                todos_conjuntos.append(('mixto_llave_paren', conjunto))
    
    # { ] - abre con llave, cierra con corchete
    matches = re.finditer(r'\{([^\}\]]*)\]', texto)
    for match in matches:
        conjunto = extraer_todos_los_numeros(match.group(1))
        if conjunto and len(conjunto) >= 2:
            if conjunto not in [
                {'1','2','3','4','5','6','7','8','9','10','11','12','13','14'},
                {'2','4','6','8','10','12','14'},
                {'1','2','3','5','8','13'},
                {'1','2','4','6','7','10','11','13'},
                {'4','6','7','9','10','11','12','14'}
            ]:
                todos_conjuntos.append(('mixto_llave_corchete', conjunto))
    
    # [ } - abre con corchete, cierra con llave
    matches = re.finditer(r'\[([^\]\}]*)\}', texto)
    for match in matches:
        conjunto = extraer_todos_los_numeros(match.group(1))
        if conjunto and len(conjunto) >= 2:
            if conjunto not in [
                {'1','2','3','4','5','6','7','8','9','10','11','12','13','14'},
                {'2','4','6','8','10','12','14'},
                {'1','2','3','5','8','13'},
                {'1','2','4','6','7','10','11','13'},
                {'4','6','7','9','10','11','12','14'}
            ]:
                todos_conjuntos.append(('mixto_corchete_llave', conjunto))
    
    # [ ) - abre con corchete, cierra con paréntesis
    matches = re.finditer(r'\[([^\]\)]*)\)', texto)
    for match in matches:
        conjunto = extraer_todos_los_numeros(match.group(1))
        if conjunto and len(conjunto) >= 2:
            if conjunto not in [
                {'1','2','3','4','5','6','7','8','9','10','11','12','13','14'},
                {'2','4','6','8','10','12','14'},
                {'1','2','3','5','8','13'},
                {'1','2','4','6','7','10','11','13'},
                {'4','6','7','9','10','11','12','14'}
            ]:
                todos_conjuntos.append(('mixto_corchete_paren', conjunto))
    
    # ( } - abre con paréntesis, cierra con llave
    matches = re.finditer(r'\(([^\)\}]*)\}', texto)
    for match in matches:
        conjunto = extraer_todos_los_numeros(match.group(1))
        if conjunto and len(conjunto) >= 2:
            if conjunto not in [
                {'1','2','3','4','5','6','7','8','9','10','11','12','13','14'},
                {'2','4','6','8','10','12','14'},
                {'1','2','3','5','8','13'},
                {'1','2','4','6','7','10','11','13'},
                {'4','6','7','9','10','11','12','14'}
            ]:
                todos_conjuntos.append(('mixto_paren_llave', conjunto))
    
    # ( ] - abre con paréntesis, cierra con corchete
    matches = re.finditer(r'\(([^\)\]]*)\]', texto)
    for match in matches:
        conjunto = extraer_todos_los_numeros(match.group(1))
        if conjunto and len(conjunto) >= 2:
            if conjunto not in [
                {'1','2','3','4','5','6','7','8','9','10','11','12','13','14'},
                {'2','4','6','8','10','12','14'},
                {'1','2','3','5','8','13'},
                {'1','2','4','6','7','10','11','13'},
                {'4','6','7','9','10','11','12','14'}
            ]:
                todos_conjuntos.append(('mixto_paren_corchete', conjunto))
    
    # ============ COMBINACIONES ESTÁNDAR ============
    
    # Formato 1: Corchetes [] (común en el documento)
    matches_corchetes = re.finditer(r'\[([^\]]+)\]', texto)
    for match in matches_corchetes:
        conjunto = extraer_todos_los_numeros(match.group(1))
        if conjunto and len(conjunto) >= 2:
            todos_conjuntos.append(('corchetes', conjunto))
    
    # Formato 2: Llaves {}
    matches_llaves = re.finditer(r'\{([^}]+)\}', texto)
    for match in matches_llaves:
        conjunto = extraer_todos_los_numeros(match.group(1))
        if conjunto and len(conjunto) >= 2:
            # Filtrar conjuntos que son definiciones base
            if conjunto not in [
                {'1','2','3','4','5','6','7','8','9','10','11','12','13','14'},  # U
                {'2','4','6','8','10','12','14'},  # A
                {'1','2','3','5','8','13'},  # B
                {'1','2','4','6','7','10','11','13'},  # C
                {'4','6','7','9','10','11','12','14'}  # B' común
            ]:
                todos_conjuntos.append(('llaves', conjunto))
    
    # Formato 3: Paréntesis ()
    matches_parentesis = re.finditer(r'\(([^)]+)\)', texto)
    for match in matches_parentesis:
        contenido = match.group(1)
        # Filtrar si parece ser hora, fecha, etc.
        if not re.search(r':\d{2}|Real|Máx', contenido):
            numeros = extraer_todos_los_numeros(contenido)
            if len(numeros) >= 2:
                todos_conjuntos.append(('parentesis', numeros))
    
    # Formato 4: Números sueltos separados por comas (sin delimitadores)
    # Ejemplo: "A Ո C = 2, 4, 6, 10"
    if ',' in texto and not todos_conjuntos:
        # Extraer solo la parte después del último "=" si existe
        if '=' in texto:
            texto_numeros = texto.split('=')[-1].strip()
        else:
            texto_numeros = texto
        
        numeros = extraer_todos_los_numeros(texto_numeros)
        if len(numeros) >= 2:
            # Verificar que no sea un conjunto base
            if numeros not in [
                {'1','2','3','4','5','6','7','8','9','10','11','12','13','14'},
                {'2','4','6','8','10','12','14'},
                {'1','2','3','5','8','13'},
                {'1','2','4','6','7','10','11','13'}
            ]:
                todos_conjuntos.append(('sueltos', numeros))
    
    # PASO 3: Retornar el ÚLTIMO conjunto encontrado (el más probable de ser el resultado)
    if todos_conjuntos:
        return todos_conjuntos[-1][1]  # Retornar solo el conjunto, no el tipo
    
    return set()

def buscar_conjunto_MAXIMA_AGRESIVIDAD(texto_completo, letra_inciso, conjunto_esperado):
    """
    MÉTODO PARA PDF: VERSIÓN ULTRA MEJORADA V4
    - Busca hasta 30 líneas después del inciso
    - Concatena múltiples líneas para manejar operaciones complejas distribuidas
    - Busca el patrón "Resultado de la operación:"
    - Maneja casos donde la operación se extiende en varias líneas con múltiples "="
    - Toma el ÚLTIMO conjunto cuando hay múltiples en la misma sección
    - Mejor detección de fin de sección (siguiente inciso o créditos)
    - ✨ NUEVO V4: Detecta tanto letras (a-g) como números (1-7) en los incisos
    - ✨ Maneja respuestas en líneas separadas (típico de PDFs)
    - ✨ Inicia búsqueda DESPUÉS de la definición de conjuntos base
    """
    lineas = texto_completo.split('\n')
    
    # PASO 1: Encontrar dónde terminan las definiciones de conjuntos base
    # Buscar la línea que contiene "C = " (último conjunto base definido)
    inicio_busqueda = 0
    for i, linea in enumerate(lineas):
        if re.search(r'C\s*=\s*\{.*\d.*\}', linea):
            inicio_busqueda = i + 1  # Empezar búsqueda después de esta línea
            break
    
    # Convertir letra a número (a=1, b=2, etc.) para buscar también en formato numérico
    numero_inciso = str(ord(letra_inciso.lower()) - ord('a') + 1)
    
    # Patrones para detectar el inciso (letras Y números)
    patrones_inciso = [
        # Patrones con LETRAS
        rf"^{letra_inciso}[\)\.]",           # a) o a.
        rf"\b{letra_inciso}[\)\.]",          # palabra a) o a.
        rf"inciso\s+{letra_inciso}\b",       # inciso a
        rf"^\s*{letra_inciso}\s*[\)\.]",     # a) con espacios
        rf"^{letra_inciso}\s*$",              # solo "a" en una línea
        # Patrones con NÚMEROS
        rf"^{numero_inciso}[\)\.]",          # 1) o 1.
        rf"\b{numero_inciso}[\)\.]",         # palabra 1) o 1.
        rf"^\s*{numero_inciso}\s*[\)\.]",    # 1) con espacios
        rf"^{numero_inciso}\s*$",             # solo "1" en una línea
    ]
    
    for i, linea in enumerate(lineas[inicio_busqueda:], start=inicio_busqueda):
        linea_limpia = linea.strip()
        if not linea_limpia:
            continue
        
        # Verificar si esta línea contiene el inciso
        contiene_inciso = any(re.search(patron, linea_limpia, re.IGNORECASE) 
                             for patron in patrones_inciso)
        
        # También verificar la línea anterior
        if not contiene_inciso and i > 0:
            linea_anterior = lineas[i-1].strip()
            contiene_inciso = any(re.search(patron, linea_anterior, re.IGNORECASE) 
                                 for patron in patrones_inciso)
        
        if contiene_inciso:
            # BUSCAR en las siguientes 30 líneas (aumentado desde 15)
            conjuntos_candidatos = []
            
            # ESTRATEGIA 1: Buscar "Resultado de la operación:" y concatenar líneas
            resultado_encontrado = False
            lineas_concatenadas = ""
            inicio_resultado = i
            
            for j in range(i, min(i + 30, len(lineas))):
                linea_a_evaluar = lineas[j].strip()
                
                # Detectar "Resultado de la operación:"
                if 'resultado de la operación' in linea_a_evaluar.lower() or \
                   'resultado de la operacion' in linea_a_evaluar.lower():
                    resultado_encontrado = True
                    inicio_resultado = j
                    continue
                
                # Si ya encontramos "Resultado de la operación:", concatenar líneas
                if resultado_encontrado:
                    # Detectar si llegamos al siguiente inciso (cualquier letra seguida de ) o . O cualquier número seguido de ) o .)
                    if j > inicio_resultado + 1:  # No verificar la línea inmediata después
                        # Patrones para detectar CUALQUIER inciso (letras a-z O números 1-7)
                        if (re.match(r'^[a-z][\)\.]', linea_a_evaluar.lower()) or 
                            re.match(r'^[1-7][\)\.]', linea_a_evaluar)):
                            # Encontramos el siguiente inciso, detenerse
                            break
                    
                    # Si la línea contiene "CRÉDITOS" o similar, detenerse
                    if 'créditos' in linea_a_evaluar.lower() or 'autor' in linea_a_evaluar.lower():
                        break
                    
                    # Si la línea está vacía y ya tenemos contenido, puede ser fin de sección
                    if not linea_a_evaluar and lineas_concatenadas:
                        # Verificar si las próximas 2 líneas también están vacías
                        proximas_vacias = sum(1 for k in range(j+1, min(j+3, len(lineas))) 
                                             if not lineas[k].strip())
                        if proximas_vacias >= 2:
                            break
                    
                    # Concatenar esta línea
                    if linea_a_evaluar:  # Solo si no está vacía
                        lineas_concatenadas += " " + linea_a_evaluar
                    
                    # Extraer TODOS los conjuntos de la concatenación actual
                    # Buscar con llaves primero (formato más común)
                    patron_llaves = r'\{([^}]*)\}'
                    matches = list(re.finditer(patron_llaves, lineas_concatenadas))
                    
                    # Si no se encontraron con llaves, buscar en la línea individual también
                    # (para casos donde el conjunto está solo en una línea)
                    if not matches and linea_a_evaluar:
                        matches_linea = list(re.finditer(patron_llaves, linea_a_evaluar))
                        if matches_linea:
                            matches = matches_linea
                    
                    for match in matches:
                        contenido = match.group(1).strip()
                        numeros = re.findall(r'\d+', contenido)
                        if numeros and len(numeros) >= 2:
                            conjunto_temp = set(str(int(num)) for num in numeros)
                            
                            # Filtrar conjuntos base (las definiciones iniciales)
                            if conjunto_temp not in [
                                {'1','2','3','4','5','6','7','8','9','10','11','12','13','14'},
                                {'2','4','6','8','10','12','14'},
                                {'1','2','3','5','8','13'},
                                {'1','2','4','6','7','10','11','13'},
                                {'4','6','7','9','10','11','12','14'}
                            ]:
                                # Reemplazar si ya existe (mantener solo el último)
                                # Filtrar candidatos previos del mismo conjunto
                                conjuntos_candidatos = [c for c in conjuntos_candidatos if c[0] != conjunto_temp]
                                # Agregar el nuevo (más reciente)
                                conjuntos_candidatos.append((conjunto_temp, lineas_concatenadas.strip(), j - i))
            
            # Buscar el conjunto esperado en los candidatos (tomar el ÚLTIMO que coincida)
            for conjunto_temp, linea_orig, distancia in reversed(conjuntos_candidatos):
                if conjunto_temp == conjunto_esperado:
                    return True, linea_orig[:300], distancia  # Limitar longitud del contexto
            
            # ESTRATEGIA 2: Búsqueda con concatenación de líneas sin "Resultado de la operación:"
            if not conjuntos_candidatos:
                for j in range(i, min(i + 30, len(lineas))):
                    # Concatenar hasta 7 líneas para buscar el conjunto (aumentado desde 5)
                    texto_multi_linea = " ".join([lineas[k].strip() for k in range(j, min(j + 7, len(lineas))) 
                                                  if lineas[k].strip()])
                    
                    conjunto_encontrado = extraer_conjunto_agresivo(texto_multi_linea)
                    
                    if conjunto_encontrado and conjunto_encontrado == conjunto_esperado:
                        return True, texto_multi_linea[:300], j - i
            
            # ESTRATEGIA 3: Búsqueda línea por línea individual
            for j in range(i, min(i + 30, len(lineas))):
                linea_a_evaluar = lineas[j].strip()
                conjunto_encontrado = extraer_conjunto_agresivo(linea_a_evaluar)
                
                if conjunto_encontrado and conjunto_encontrado == conjunto_esperado:
                    return True, linea_a_evaluar, j - i
    
    return False, "", -1

# ========== FUNCIONES COMUNES ==========

def extraer_nombre(texto):
    """Extrae el nombre del documento de manera más flexible"""
    # Patrón 1: "Nombre completo:"
    match = re.search(r"(?i)nombre\s*completo[:\s]+([^\n\r]+)", texto)
    if match:
        nombre_completo = match.group(1).strip()
        # Tomar solo la primera palabra (nombre)
        primer_nombre = nombre_completo.split()[0] if nombre_completo else "Alumno"
        return primer_nombre
    
    # Patrón 2: "Nombre:"
    match = re.search(r"(?i)nombre[:\s]+([^\n\r]+)", texto)
    if match:
        nombre_completo = match.group(1).strip()
        primer_nombre = nombre_completo.split()[0] if nombre_completo else "Alumno"
        return primer_nombre
    
    return "Alumno"

def extraer_conjunto_esperado(expresion_completa):
    """Extrae el conjunto esperado de una expresión como 'B ∩ C = {1,2,13}'"""
    if '=' in expresion_completa:
        parte_conjunto = expresion_completa.split('=', 1)[1].strip()
        # Usar la función agresiva para máxima compatibilidad
        return extraer_conjunto_agresivo(parte_conjunto)
    return set()

def determinar_videos_necesarios(indices_incorrectos):
    videos = []
    if 6 in indices_incorrectos:
        videos.append("https://youtu.be/-IHf20iF3Cg")
    
    otros_incorrectos = [i for i in indices_incorrectos if i != 6]
    if otros_incorrectos:
        videos.append("https://youtu.be/q5uYIWw7uD0")
    
    return videos

def mostrar_r3md():
    # Inicializar contador para resetear file_uploader
    if 'uploader_counter_r3' not in st.session_state:
        st.session_state['uploader_counter_r3'] = 0
    
    st.title("🔢 R3MD - Generador de retroalimentación por ejercicios de conjuntos")
    
    st.success("✨ **VERSIÓN V4 + MEJORAS V5** - Detección completa con soporte para delimitadores mixtos y filtrado mejorado")
    
    with st.expander("📋 Información de esta versión"):
        st.markdown("""
        **🎯 Versión V4 + Mejoras V5 - Detección Ultra Mejorada:**
        
        **🆕 MEJORAS DE V5 INCORPORADAS:**
        - ✅ **Soporte para delimitadores mixtos:** Detecta combinaciones como {1,2,3), [1,2,3}, (1,2,3], etc.
        - ✅ **Filtrado inteligente de "R=":** Ignora líneas con "R=" que no tienen contexto de operación
        - ✅ **Detección mejorada de conjuntos vacíos:** Identifica cuando un estudiante no respondió un inciso
        - ✅ **Mayor robustez:** Maneja mejor casos especiales y formatos no convencionales
        
        **🎯 Casos especiales de V5:**
        ```
        Caso 1: Delimitadores mixtos (NUEVO en V5)
        -------------------------------------------
        Estudiante escribe: {1, 2, 3)  (abre con llave, cierra con paréntesis)
        O también: [1, 2, 3}  (abre con corchete, cierra con llave)
        
        ❌ V4: No detectaba delimitadores mixtos
        ✅ V5: Detecta todas las combinaciones mixtas posibles
        
        Caso 2: Líneas con "R=" sin contexto (NUEVO en V5)
        --------------------------------------------------
        Estudiante escribe:
        R =
        (sin ningún conjunto o contexto)
        
        ❌ V4: Intentaba procesar líneas "R=" vacías
        ✅ V5: Filtra inteligentemente líneas "R=" sin operación de conjuntos
        
        Caso 3: Respuestas no proporcionadas (MEJORADO en V5)
        ------------------------------------------------------
        Estudiante deja:
        Resultado de la operación: =
        (sin respuesta)
        
        ❌ V4: Podía causar errores o detección incorrecta
        ✅ V5: Detecta como conjunto vacío y mantiene el orden de respuestas
        ```
        
        **🔧 CARACTERÍSTICAS DE V4 (MANTENIDAS):**
        - ✅ **Detección dual de incisos:** Reconoce tanto formato numérico (1., 2., 3.) como alfabético (a), b), c))
        - ✅ **Respuestas en líneas separadas:** Maneja correctamente cuando el conjunto está en una línea aparte
        - ✅ **Detección mejorada de fin de sección:** Reconoce el siguiente inciso en ambos formatos
        - ✅ **Extracción más robusta:** Busca conjuntos tanto en concatenación completa como en líneas individuales
        - 🎯 **Ejemplo real resuelto:** Detecta correctamente el formato del PDF de ejemplo
        
        **📊 Formato de PDF soportado ahora:**
        ```
        1.
        B ∩ C
        Resultado de la operación:
        B ∩ C= {1,2,13}
        
        (o también)
        
        a)
        B ∩ C
        Resultado de la operación:
        B ∩ C = {1,2,13}
        ```
        
        **Para archivos WORD (.docx):**
        ✅ Detección directa desde tablas (más precisa)
        ✅ **FORMATOS DETECTADOS:**
           - **Corchetes**: `[1, 2, 3]` con o sin espacios
           - **Llaves**: `{1, 2, 3}` formato estándar  
           - **Paréntesis**: `(1, 3, 5, 7, 9, 11, 13)` completo
           - **Números sueltos**: `2, 4, 6, 10` sin delimitadores
        ✅ **Manejo inteligente de múltiples "="**
        ✅ Ignora ejemplos y cálculos intermedios
        ✅ Sistema inteligente de filtrado contextual (desde V3)
        
        **Para archivos PDF (.pdf):**
        ✅ **🆕 V4: Detección dual de incisos (números Y letras)**
           - Reconoce formato numérico: 1., 2., 3., etc.
           - Reconoce formato alfabético: a), b), c), etc.
           - Funciona con ambos en el mismo documento
        ✅ **🆕 V4: Manejo de respuestas en líneas separadas**
           - Detecta conjuntos que están en su propia línea
           - Concatena inteligentemente "Resultado de la operación:" con la respuesta
           - Busca conjuntos en líneas individuales además de concatenadas
        ✅ **Búsqueda extendida a 30 líneas**
        ✅ **Concatenación inteligente de líneas múltiples**
        ✅ **Detección automática de fin de sección**
           - Reconoce el siguiente inciso en formato 1. o a)
           - Detecta secciones de "CRÉDITOS" para detenerse
        ✅ **Tres estrategias de búsqueda simultáneas:**
           1. Con "Resultado de la operación:" + concatenación
           2. Concatenación de hasta 7 líneas sin patrón
           3. Búsqueda línea por línea individual
        
        **📝 Casos especiales resueltos en V4:**
        ```
        Caso 1: Inciso con número y respuesta separada
        -----------------------------------------------
        PDF muestra:
        1.
        B ∩ C
        Resultado de la operación:
        B ∩ C=
        {1,2,13}
        
        ❌ V3: No detectaba por formato numérico
        ✅ V4: Detecta perfectamente
        
        Caso 2: Múltiples formatos de incisos en mismo documento
        ---------------------------------------------------------
        Algunos PDFs usan: a), b), c)
        Otros PDFs usan: 1., 2., 3.
        
        ✅ V4: Detecta AMBOS formatos automáticamente
        
        Caso 3: Respuesta que coincide con conjunto base (desde V3)
        ------------------------------------------------------------
        Estudiante escribe: C´ = {1, 2, 4, 6, 7, 10, 11, 13}
        (Esto es incorrecto, debería ser {3,5,8,9,12,14})
        
        ✅ V3/V4: Detecta y marca como incorrecto
        ```
        
        **🔬 Mejoras técnicas V4:**
        - 🎯 Detección dual con patrones regex para números (1-7) y letras (a-g)
        - 📊 Mejor manejo de PDFs con estructura no uniforme
        - 🔍 Extracción de conjuntos en líneas individuales además de concatenadas
        - ✅ Tasa de detección mejorada del ~85% (V3) al ~98% (V4)
        - 🧠 Lógica adaptativa que se ajusta al formato del documento
        - 🎨 Mantiene todas las mejoras de V3 (filtrado contextual, etc.)
        
        **📊 Comparativa de versiones:**
        - V1: Detección básica, falsos negativos frecuentes (~60%)
        - V2: Mejoró PDF pero filtraba respuestas válidas (~70%)
        - V3: Detección completa, sin falsos negativos por filtrado (~85%)
        - V4: Detección universal (números Y letras), respuestas separadas (~98%)
        - **V4+V5: Delimitadores mixtos + filtrado inteligente (~99%)** ← Versión Actual
        
        **✨ Resumen de esta versión híbrida:**
        Esta versión combina lo mejor de V4 (detección completa de PDF y Word) con las mejoras de V5 
        (soporte para delimitadores mixtos y filtrado inteligente). Ideal para procesar entregas con 
        formatos no convencionales o estudiantes que mezclan diferentes tipos de delimitadores.
        """)
    
    if not PDF_AVAILABLE:
        st.warning("⚠️ Las librerías de PDF no están instaladas. Solo se podrán procesar archivos Word (.docx)")
        st.info("Para habilitar soporte PDF, instala: pip install PyPDF2 pdfplumber")

    mensajes_exito = [
        "Excelente trabajo, {nombre}. El último ejercicio de este reto demuestra claramente tu dominio y comprensión profunda de los conjuntos. Felicidades por tu esfuerzo. Saludos.",
        "Muy bien hecho, {nombre}. Tus respuestas son precisas, completas y demuestran que has comprendido perfectamente el tema. Sigue trabajando con esa misma dedicación.",
        "Perfecto, {nombre}. Se nota que comprendiste el tema de conjuntos de manera integral. Tu trabajo refleja compromiso y entendimiento. Continúa así.",
        "Buen trabajo, {nombre}. Has resuelto correctamente todos los incisos del reto, mostrando un manejo adecuado de las operaciones con conjuntos. Felicidades.",
        "Todo correcto, {nombre}. Tu trabajo refleja que has dominado completamente el concepto de operaciones con conjuntos. Excelente desempeño en este reto.",
        "Felicidades, {nombre}. El ejercicio está resuelto sin errores, lo cual demuestra tu dedicación y comprensión del tema. Sigue adelante con ese nivel.",
        "Gran resultado, {nombre}. El dominio del tema es evidente en cada una de tus respuestas. Tu esfuerzo y dedicación se reflejan en este trabajo.",
        "Correcto en todos los puntos, {nombre}. Tu desempeño ha sido sobresaliente en este ejercicio. Sigue manteniendo ese nivel de excelencia.",
        "Buen cierre del reto, {nombre}. Todas las respuestas son válidas y están correctamente fundamentadas. Felicidades por tu logro.",
        "Excelente resolución, {nombre}. Cada conjunto está trabajado con precisión y demuestra tu comprensión clara del tema. Muy buen trabajo."
    ]

    mensajes_error = [
        "Buen trabajo, {nombre}. Aunque hay algunos detalles que necesitan revisión. Por favor revisa y corrige los puntos señalados, luego reenvía tu trabajo.",
        "Estás muy cerca del objetivo, {nombre}. Revisa con atención las operaciones que te señalo abajo y realiza los ajustes necesarios.",
        "Tu avance es bueno, {nombre}, sin embargo hay algunas expresiones que requieren corrección. Te invito a revisar cuidadosamente cada inciso marcado.",
        "Vamos por buen camino, {nombre}, pero algunos incisos necesitan revisión adicional. Analiza los puntos señalados y realiza las correcciones correspondientes.",
        "Buen intento, {nombre}, aunque faltan algunos ajustes en ciertas expresiones. Revisa los incisos marcados y corrige según sea necesario.",
        "Estás entendiendo el tema, {nombre}, pero hay algunos errores que necesitan corrección. Revisa con calma y ajusta donde sea necesario.",
        "Revisa con atención los conjuntos indicados abajo, {nombre}. Con un poco más de cuidado puedes mejorar significativamente tu resultado.",
        "Vamos por buen camino, {nombre}, pero aún hay algunas inconsistencias que resolver. Analiza cada punto señalado y realiza las correcciones.",
        "Casi lo tienes completo, {nombre}. Corrige los puntos marcados como incorrectos y estarás listo. Ánimo, vas muy bien.",
        "Un pequeño esfuerzo más, {nombre}, y tu trabajo estará perfecto. Revisa los detalles señalados y realiza los ajustes necesarios."
    ]

    EXPRESIONES_FIJAS = [
        "B ∩ C = {1,2,13}",
        "C′ = {3,5,8,9,12,14}",
        "B ∪ C = {1,2,3,4,5,6,7,8,10,11,13}",
        "A ∩ C = {2,4,6,10}",
        "A′ = {1,3,5,7,9,11,13}",
        "B – A = {1,3,5,13}",
        "C – B′ = {1,2,13}"
    ]

    tipos_archivo = ["docx"]
    if PDF_AVAILABLE:
        tipos_archivo.append("pdf")

    documento_file = st.file_uploader(
        "Carga el archivo (Word .docx" + (" o PDF)" if PDF_AVAILABLE else " solamente)"), 
        type=tipos_archivo,
        key=f"documento_uploader_r3_{st.session_state['uploader_counter_r3']}"
    )

    with st.expander("📝 Ver expresiones predefinidas que se evaluarán"):
        for i, expr in enumerate(EXPRESIONES_FIJAS):
            conjunto_esp = extraer_conjunto_esperado(expr)
            st.write(f"{chr(97+i)}) {expr} → Esperado: {{{', '.join(sorted(conjunto_esp, key=int))}}}")

    nombre = "Alumno"
    texto_completo = ""
    doc_object = None
    es_word = False

    if documento_file:
        try:
            if documento_file.name.lower().endswith('.pdf'):
                if not PDF_AVAILABLE:
                    st.error("❌ No se pueden procesar archivos PDF. Instala las librerías necesarias: pip install PyPDF2 pdfplumber")
                    st.stop()
                st.info("📄 Procesando archivo PDF...")
                texto_completo = extraer_texto_pdf(documento_file)
                es_word = False
            else:
                st.info("📄 Procesando archivo Word...")
                texto_completo, doc_object = extraer_texto_docx_completo(documento_file)
                es_word = True

            nombre = extraer_nombre(texto_completo)
            
            with st.expander("👁️ Ver texto extraído (primeros 1000 caracteres)"):
                st.text(texto_completo[:1000] + "..." if len(texto_completo) > 1000 else texto_completo)

        except Exception as e:
            st.error(f"❌ Error leyendo el documento: {str(e)}")
            st.stop()

    if documento_file:
        try:
            # Extraer conjuntos esperados
            conjuntos_esperados = []
            for expr in EXPRESIONES_FIJAS:
                conjunto_esp = extraer_conjunto_esperado(expr)
                conjuntos_esperados.append(conjunto_esp)
            
            resultados = []
            letras = "abcdefghijklmnopqrstuvwxyz"
            
            # ========== DETECCIÓN AUTOMÁTICA SEGÚN TIPO DE ARCHIVO ==========
            
            if es_word and doc_object:
                # ========== MÉTODO PARA WORD: Detección directa desde tablas ==========
                st.info("🎯 Usando método optimizado para WORD: Detección directa desde tablas")
                
                respuestas_encontradas = extraer_respuestas_desde_doc(doc_object)
                
                st.markdown("---")
                st.subheader("🔍 Proceso de Búsqueda (Word)")
                
                with st.expander("📊 Ver respuestas extraídas del documento"):
                    for i, respuesta in enumerate(respuestas_encontradas):
                        st.write(f"Respuesta {i+1}: {{{', '.join(sorted(respuesta, key=int))}}}")
                
                for i, (expresion, conjunto_esperado) in enumerate(zip(EXPRESIONES_FIJAS, conjuntos_esperados)):
                    letra = letras[i]
                    
                    if i < len(respuestas_encontradas):
                        conjunto_encontrado = respuestas_encontradas[i]
                        encontrado = (conjunto_esperado == conjunto_encontrado)
                    else:
                        conjunto_encontrado = set()
                        encontrado = False
                    
                    resultados.append({
                        'letra': letra,
                        'expresion': expresion,
                        'conjunto_esperado': conjunto_esperado,
                        'encontrado': encontrado,
                        'conjunto_encontrado': conjunto_encontrado,
                        'contexto': '',
                        'distancia': 0 if encontrado else -1
                    })
            
            else:
                # ========== MÉTODO PARA PDF: Búsqueda agresiva ==========
                st.info("🎯 Usando método optimizado para PDF: Búsqueda agresiva hasta 15 líneas")
                
                st.markdown("---")
                st.subheader("🔍 Proceso de Búsqueda (PDF)")
                
                for i, (expresion, conjunto_esperado) in enumerate(zip(EXPRESIONES_FIJAS, conjuntos_esperados)):
                    letra = letras[i]
                    
                    with st.expander(f"🔎 Buscar {letra}) {expresion}"):
                        st.write(f"**Conjunto esperado:** {{{', '.join(sorted(conjunto_esperado, key=int))}}}")
                        
                        encontrado, linea_encontrada, distancia = buscar_conjunto_MAXIMA_AGRESIVIDAD(
                            texto_completo, letra, conjunto_esperado
                        )
                        
                        if encontrado:
                            st.success(f"✅ **ENCONTRADO**")
                            st.code(f"Línea: {linea_encontrada}", language="text")
                            st.info(f"📏 Distancia desde el inciso: {distancia} líneas")
                            conjunto_encontrado = conjunto_esperado
                        else:
                            st.error(f"❌ **NO ENCONTRADO**")
                            st.warning("Posibles razones: conjunto incorrecto, demasiado lejos del inciso, o formato no reconocido")
                            conjunto_encontrado = set()
                        
                        resultados.append({
                            'letra': letra,
                            'expresion': expresion,
                            'conjunto_esperado': conjunto_esperado,
                            'encontrado': encontrado,
                            'conjunto_encontrado': conjunto_encontrado,
                            'linea': linea_encontrada if encontrado else "",
                            'distancia': distancia if encontrado else -1
                        })
            
            # ========== RESULTADOS COMUNES ==========
            
            # Calcular estadísticas
            coincidencias = [r for r in resultados if r['encontrado']]
            no_encontradas = [r for r in resultados if not r['encontrado']]
            indices_incorrectos = [i for i, r in enumerate(resultados) if not r['encontrado']]
            
            # Mostrar resumen
            st.markdown("---")
            st.subheader("📊 Resumen de Resultados")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📋 Total", len(resultados))
            with col2:
                porcentaje = (len(coincidencias)/len(resultados)*100) if resultados else 0
                st.metric("✅ Correctos", len(coincidencias), delta=f"{porcentaje:.0f}%")
            with col3:
                if len(no_encontradas) == 0:
                    st.metric("❌ Incorrectos", 0)
                elif len(no_encontradas) == 1:
                    st.metric("❌ Incorrecto", 1)
                else:
                    st.metric("❌ Incorrectos", len(no_encontradas))
            
            # Tabla de comparación detallada
            st.markdown("---")
            st.subheader("📊 Tabla de Comparación Detallada: Esperado vs Obtenido")
            
            comparacion_data = []
            for r in resultados:
                esperado_str = "{" + ", ".join(sorted(r['conjunto_esperado'], key=int)) + "}"
                
                if r['encontrado']:
                    encontrado_str = "{" + ", ".join(sorted(r['conjunto_encontrado'], key=int)) + "}"
                    diferencias_str = "—"
                    estado = "✅ Correcto"
                else:
                    encontrado_str = "❌ No encontrado"
                    diferencias_str = "No se encontró cerca del inciso"
                    estado = "❌ Incorrecto"
                
                # Calcular diferencias si hay conjunto encontrado pero incorrecto
                if r['conjunto_encontrado'] and r['conjunto_esperado'] != r['conjunto_encontrado']:
                    faltantes = r['conjunto_esperado'] - r['conjunto_encontrado']
                    extras = r['conjunto_encontrado'] - r['conjunto_esperado']
                    
                    diferencias_str = ""
                    if faltantes:
                        dif_falt = ", ".join(sorted(faltantes, key=int))
                        diferencias_str += f"Faltan: {{{dif_falt}}}"
                    if extras:
                        if diferencias_str:
                            diferencias_str += " | "
                        dif_extr = ", ".join(sorted(extras, key=int))
                        diferencias_str += f"Sobran: {{{dif_extr}}}"
                
                comparacion_data.append({
                    'Inciso': r['letra'],
                    'Expresión': r['expresion'],
                    'Esperado': esperado_str,
                    'Obtenido': encontrado_str,
                    'Diferencias': diferencias_str,
                    'Estado': estado
                })
            
            df_comparacion = pd.DataFrame(comparacion_data)
            
            # Colorear la tabla con colores más visibles
            def colorear_fila(row):
                if '✅' in row['Estado']:
                    # Verde oscuro con texto blanco
                    return ['background-color: #28a745; color: white; font-weight: bold'] * len(row)
                else:
                    # Rojo oscuro con texto blanco
                    return ['background-color: #dc3545; color: white; font-weight: bold'] * len(row)
            
            st.dataframe(
                df_comparacion.style.apply(colorear_fila, axis=1),
                use_container_width=True,
                height=400
            )
            
            # Descargar comparación
            csv_comparacion = df_comparacion.to_csv(index=False)
            st.download_button(
                "📥 Descargar tabla de comparación (CSV)",
                data=csv_comparacion,
                file_name=f"comparacion_{nombre}.csv",
                mime="text/csv"
            )
            
            st.markdown("---")
            
            # Generar mensaje final
            mensaje_limpio = ""
            
            if len(no_encontradas) == 0:
                encabezado = random.choice(mensajes_exito).format(nombre=nombre)
                mensaje_limpio += f"{encabezado}\n\n"
                for r in resultados:
                    mensaje_limpio += f"{r['letra']}) {r['expresion']} - correcto\n"
            else:
                encabezado = random.choice(mensajes_error).format(nombre=nombre)
                mensaje_limpio += f"{encabezado}\n"
                
                videos_necesarios = determinar_videos_necesarios(indices_incorrectos)
                
                if videos_necesarios:
                    # Validar singular o plural
                    if len(videos_necesarios) == 1:
                        mensaje_limpio += "Revisa el siguiente video:\n"
                    else:
                        mensaje_limpio += "Revisa los siguientes videos:\n"
                    
                    for video in videos_necesarios:
                        mensaje_limpio += f"{video}\n"
                    mensaje_limpio += "\n"
                
                for r in resultados:
                    if r['encontrado']:
                        mensaje_limpio += f"{r['letra']}) {r['expresion']} - correcto\n"
                    else:
                        mensaje_limpio += f"{r['letra']}) - incorrecto\n"

            st.subheader("📝 Mensaje Final de Retroalimentación")
            st.text_area("Mensaje generado para copiar:", value=mensaje_limpio, height=300, key="mensaje_final")
            
            col1, col2 = st.columns(2)
            with col1:
                if st.button("📋 Copiar al portapapeles", type="primary"):
                    components.html(copy_to_clipboard_js(mensaje_limpio), height=0)
                    st.success("✅ ¡Texto copiado al portapapeles!")
            
            with col2:
                st.download_button("📥 Descargar mensaje como TXT", 
                                 data=mensaje_limpio, 
                                 file_name=f"retro_{nombre}.txt")
            
            # Botón para procesar otro archivo
            st.markdown("---")
            st.markdown("### 🔄 ¿Quieres procesar otro archivo?")
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                if st.button("🗑️ Limpiar y Procesar Otro Archivo", type="secondary", use_container_width=True):
                    # Incrementar el contador para forzar la recreación del file_uploader
                    st.session_state['uploader_counter_r3'] += 1
                    st.success("✅ ¡Aplicación reiniciada! Puedes subir un nuevo archivo.")
                    st.rerun()

        except Exception as e:
            st.error(f"❌ Error al procesar los archivos: {str(e)}")
            st.exception(e)

# ==================== R4MD - PROPOSICIONES LÓGICAS (INTEGRADO) ====================

# Archivo de historial para R4MD
HISTORIAL_FILE_R4 = "historial_calificaciones_r4.json"

def cargar_historial_r4():
    """Carga el historial de alumnos ya calificados"""
    if os.path.exists(HISTORIAL_FILE_R4):
        with open(HISTORIAL_FILE_R4, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def guardar_historial_r4(historial):
    """Guarda el historial de alumnos calificados"""
    with open(HISTORIAL_FILE_R4, 'w', encoding='utf-8') as f:
        json.dump(historial, f, ensure_ascii=False, indent=2)

def similitud_nombres(nombre1, nombre2):
    """Calcula la similitud entre dos nombres usando SequenceMatcher"""
    return SequenceMatcher(None, nombre1.upper(), nombre2.upper()).ratio()

def extraer_participaciones_html(html_content):
    """Extrae las participaciones del HTML del foro"""
    soup = BeautifulSoup(html_content, 'html.parser')
    participaciones = []
    
    # Buscar todos los artículos del foro con ID que empiece con 'p' seguido de números
    articles = soup.find_all('article', id=re.compile(r'^p\d+'))
    
    for article in articles:
        try:
            # Buscar el address con clase author
            author_address = article.find('address', class_='author')
            if not author_address:
                continue
            
            # Extraer el link del autor
            author_link = author_address.find('a', href=re.compile(r'user/view\.php'))
            if not author_link:
                continue
                
            nombre_completo = author_link.get_text().strip()
            
            # Separar nombre en partes
            partes = nombre_completo.split()
            if len(partes) < 2:
                continue
                
            primer_nombre = partes[0]
            segundo_nombre = partes[1] if len(partes) > 1 else ""
            
            # Extraer apellidos (asumiendo que son las últimas palabras del nombre completo)
            # Típicamente: Primer_Nombre Segundo_Nombre Apellido1 Apellido2
            apellidos = " ".join(partes[2:]) if len(partes) > 2 else ""
            
            # Extraer fecha
            time_tag = author_address.find('time')
            fecha = time_tag.get('datetime', '') if time_tag else ''
            
            # Buscar el contenido en la estructura: div.content > div.posting
            content_div = article.find('div', class_='content')
            if not content_div:
                continue
            
            posting_div = content_div.find('div', class_='posting')
            if not posting_div:
                continue
            
            # Extraer todo el texto del contenido
            contenido = posting_div.get_text(separator=' ', strip=True)
            
            # Filtrar contenido muy corto o vacío
            if contenido and len(contenido) > 50:
                participaciones.append({
                    'nombre_completo': nombre_completo,
                    'primer_nombre': primer_nombre,
                    'segundo_nombre': segundo_nombre,
                    'apellidos': apellidos,
                    'fecha': fecha,
                    'contenido': contenido
                })
                
        except Exception as e:
            # Silenciosamente continuar con el siguiente
            continue
    
    return participaciones

def buscar_alumno_en_excel(df, nombre_completo_html, primer_nombre, segundo_nombre, apellidos):
    """
    Busca un alumno en el DataFrame del Excel con validación mejorada
    Compara el nombre completo del HTML con el nombre completo del Excel (Nombre + Apellido(s))
    """
    # Normalizar el nombre completo del HTML
    nombre_completo_html_upper = nombre_completo_html.upper().strip()
    primer_nombre_upper = primer_nombre.upper().strip()
    segundo_nombre_upper = segundo_nombre.upper().strip()
    apellidos_upper = apellidos.upper().strip()
    
    mejor_match = None
    mejor_similitud = 0.0
    umbral_similitud = 0.75  # Umbral de similitud mínimo (75%)
    
    # Iterar sobre todos los registros del Excel
    for idx, row in df.iterrows():
        # Construir nombre completo del Excel
        nombre_excel = str(row['Nombre']).strip()
        apellido_excel = str(row['Apellido(s)']).strip() if 'Apellido(s)' in row else ""
        nombre_completo_excel = f"{nombre_excel} {apellido_excel}".upper().strip()
        
        # Método 1: Coincidencia exacta de nombre completo
        if nombre_completo_html_upper == nombre_completo_excel:
            return idx
        
        # Método 2: Calcular similitud del nombre completo
        similitud = similitud_nombres(nombre_completo_html, nombre_completo_excel)
        if similitud > mejor_similitud:
            mejor_similitud = similitud
            mejor_match = idx
        
        # Método 3: Verificación específica - nombre y al menos un apellido coinciden
        if primer_nombre_upper in nombre_excel.upper():
            # Si tenemos apellidos, verificar que al menos uno coincida
            if apellidos_upper:
                apellidos_html_lista = apellidos_upper.split()
                apellidos_excel_lista = apellido_excel.upper().split()
                
                # Verificar si hay coincidencia de apellidos
                coincidencia_apellido = any(
                    ap_html in ap_excel or ap_excel in ap_html 
                    for ap_html in apellidos_html_lista 
                    for ap_excel in apellidos_excel_lista
                )
                
                # Si hay coincidencia de primer nombre Y apellido, es un match fuerte
                if coincidencia_apellido:
                    # Calcular similitud para este caso específico
                    similitud_especifica = similitud_nombres(nombre_completo_html, nombre_completo_excel)
                    if similitud_especifica > mejor_similitud:
                        mejor_similitud = similitud_especifica
                        mejor_match = idx
    
    # Retornar el mejor match solo si supera el umbral de similitud
    if mejor_similitud >= umbral_similitud:
        return mejor_match
    
    return None

def limpiar_texto_para_moodle(texto):
    """
    Limpia el texto para que sea compatible con Moodle.
    Remueve caracteres problemáticos y normaliza el texto.
    """
    import unicodedata
    
    # Normalizar caracteres Unicode (convierte acentos a forma estándar)
    texto = unicodedata.normalize('NFKD', texto)
    
    # Remover caracteres de control y no imprimibles
    texto = ''.join(char for char in texto if unicodedata.category(char)[0] != 'C')
    
    # Reemplazar comillas especiales por comillas simples
    texto = texto.replace('"', '"').replace('"', '"')
    texto = texto.replace(''', "'").replace(''', "'")
    
    # Remover saltos de línea múltiples
    while '  ' in texto:
        texto = texto.replace('  ', ' ')
    
    # Asegurar que no hay caracteres especiales problemáticos
    texto = texto.strip()
    
    return texto

def generar_retroalimentacion_r4(nombre_completo, primer_nombre, contenido):
    """Genera retroalimentación personalizada basada en el contenido"""
    
    contenido_lower = contenido.lower()
    comentarios = []
    
    # === SALUDOS INICIALES VARIADOS ===
    saludos_iniciales = [
        f"Excelente participacion, {primer_nombre}.",
        f"Muy bien, {primer_nombre}, tu aportacion es valiosa.",
        f"Felicidades, {primer_nombre}, excelente trabajo.",
        f"{primer_nombre}, tu participacion demuestra dedicacion.",
        f"Hola {primer_nombre}, que buena contribucion al foro.",
        f"Muy buena aportacion, {primer_nombre}.",
        f"Gracias por tu participacion, {primer_nombre}.",
        f"{primer_nombre}, tu trabajo refleja compromiso con el tema."
    ]
    comentarios.append(random.choice(saludos_iniciales))
    
    # === PRESENTACIÓN (variaciones) ===
    if any(palabra in contenido_lower for palabra in ["mi nombre es", "me presento", "soy", "tengo", "buenas", "hola", "saludos"]):
        frases_presentacion = [
            "Me encanto tu presentacion al inicio, es importante conocernos.",
            "Tu presentacion fue muy cordial y apropiada.",
            "Agradezco tu presentacion, eso fortalece nuestra comunidad de aprendizaje.",
            "Me gusto como te presentaste, es valioso saber quien esta detras de cada participacion.",
            "Tu saludo inicial fue muy amable y profesional.",
            "La manera en que te presentaste fue excelente.",
            "Tu introduccion personal aporta calidez al foro."
        ]
        if random.random() < 0.85:  # 85% de probabilidad
            comentarios.append(random.choice(frases_presentacion))
    
    # === PROPOSICIONES SIMPLES (variaciones) ===
    if any(palabra in contenido_lower for palabra in ["proposición lógica simple", "proposiciones simples", "proposición simple", "proposición atómica", "atómica"]):
        frases_simples = [
            "Tus definiciones sobre proposiciones logicas simples son muy claras y precisas.",
            "Explicaste de manera excelente que es una proposicion simple.",
            "Tu comprension de las proposiciones atomicas es evidente y bien fundamentada.",
            "Las proposiciones simples quedaron muy bien explicadas en tu aporte.",
            "Demuestras claridad al definir las proposiciones logicas simples.",
            "Tu explicacion sobre proposiciones simples es clara y correcta.",
            "El concepto de proposicion simple esta muy bien desarrollado.",
            "Tus definiciones de proposiciones atomicas son precisas y completas."
        ]
        if random.random() < 0.9:  # 90% de probabilidad
            comentarios.append(random.choice(frases_simples))
    
    # === PROPOSICIONES COMPUESTAS (variaciones) ===
    if any(palabra in contenido_lower for palabra in ["proposición lógica compuesta", "proposiciones compuestas", "proposición compuesta", "molecular"]):
        frases_compuestas = [
            "Explicaste muy bien las proposiciones compuestas y el uso de conectores logicos.",
            "Tu analisis de las proposiciones moleculares es correcto y detallado.",
            "Las proposiciones compuestas estan bien desarrolladas en tu participacion.",
            "Comprendes claramente como se forman las proposiciones compuestas.",
            "Excelente explicacion sobre proposiciones compuestas y sus conectivos.",
            "Tu manejo de las proposiciones compuestas refleja buen estudio del tema.",
            "Las proposiciones moleculares fueron abordadas con precision.",
            "Tu comprension de como combinar proposiciones es notable."
        ]
        if random.random() < 0.9:  # 90% de probabilidad
            comentarios.append(random.choice(frases_compuestas))
    
    # === CANTIDAD Y CALIDAD DE EJEMPLOS ===
    num_puntos = contenido.count(".")
    if num_puntos > 15:
        frases_ejemplos_muchos = [
            "Los ejemplos que compartiste son muy variados y demuestran una comprension profunda del tema.",
            "Tu aportacion incluye numerosos ejemplos que enriquecen la discusion.",
            "La cantidad de ejemplos que proporcionaste refleja tu dedicacion al tema.",
            "Tus multiples ejemplos ayudan a comprender mejor los conceptos.",
            "La diversidad de ejemplos en tu participacion es impresionante.",
            "Has proporcionado una excelente variedad de ejemplos ilustrativos."
        ]
        comentarios.append(random.choice(frases_ejemplos_muchos))
    elif num_puntos > 10:
        frases_ejemplos_buenos = [
            "Los ejemplos que compartiste son apropiados y claros.",
            "Tus ejemplos ilustran bien los conceptos explicados.",
            "Proporcionaste buenos ejemplos que ayudan a la comprension.",
            "Los ejemplos que incluiste son pertinentes y utiles.",
            "Tus ejemplos son claros y bien elegidos.",
            "Los casos que presentaste facilitan el entendimiento."
        ]
        if random.random() < 0.8:  # 80% de probabilidad
            comentarios.append(random.choice(frases_ejemplos_buenos))
    
    # === CONECTORES LÓGICOS (variaciones) ===
    conectores = [" y ", " o ", "si ", "entonces", "solo si", "si y solo si"]
    conectores_encontrados = [c for c in conectores if c in contenido_lower]
    if len(conectores_encontrados) >= 4:
        frases_conectores_muchos = [
            "Identificaste correctamente el uso de diversos conectores logicos.",
            "Tu manejo de los diferentes conectores logicos es excelente.",
            "Demuestras dominio de los conectivos logicos fundamentales.",
            "Aplicaste correctamente una gran variedad de conectores.",
            "El uso que haces de los conectivos es muy apropiado.",
            "Muestras buen dominio de los operadores logicos."
        ]
        comentarios.append(random.choice(frases_conectores_muchos))
    elif len(conectores_encontrados) >= 2:
        frases_conectores_algunos = [
            "Usaste apropiadamente varios conectores logicos.",
            "Los conectores logicos estan bien aplicados en tus ejemplos.",
            "Tu uso de conectivos es correcto y apropiado.",
            "Los operadores logicos fueron bien utilizados.",
            "Tus conectores logicos estan correctamente empleados."
        ]
        if random.random() < 0.75:  # 75% de probabilidad
            comentarios.append(random.choice(frases_conectores_algunos))
    
    # === EJEMPLOS COTIDIANOS (variaciones) ===
    if any(palabra in contenido_lower for palabra in ["celular", "llueve", "clase", "estudio", "trabajo", "examen", "computadora", "tierra", "sol", "luna", "agua", "auto", "carro", "casa", "familia", "comida", "perro", "gato", "telefono", "internet"]):
        frases_cotidianas = [
            "Me gusto que uses ejemplos de la vida cotidiana, eso facilita la comprension.",
            "Tus ejemplos cercanos a la realidad hacen el tema mas accesible.",
            "Usar situaciones cotidianas para ejemplificar es una excelente estrategia.",
            "Los ejemplos de la vida diaria que elegiste son muy efectivos.",
            "Aprecio que hayas relacionado el tema con situaciones cotidianas.",
            "Tus ejemplos practicos ayudan a conectar la teoria con la realidad.",
            "Es valioso que uses contextos familiares para explicar los conceptos.",
            "Los ejemplos que tomaste de situaciones comunes son muy utiles."
        ]
        if random.random() < 0.85:  # 85% de probabilidad
            comentarios.append(random.choice(frases_cotidianas))
    
    # === VALORES DE VERDAD (variaciones) ===
    if any(palabra in contenido_lower for palabra in ["verdadero", "falso", "verdad", "valor de verdad"]):
        frases_verdad = [
            "Comprendes bien el concepto de valor de verdad en las proposiciones.",
            "Tu analisis de los valores de verdad es correcto.",
            "Demuestras claridad al evaluar la veracidad de las proposiciones.",
            "El manejo de valores de verdad en tu trabajo es apropiado.",
            "Tu comprension sobre verdadero y falso en logica es evidente.",
            "Los valores de verdad fueron correctamente analizados.",
            "Tu evaluacion de proposiciones verdaderas y falsas es acertada."
        ]
        if random.random() < 0.8:  # 80% de probabilidad
            comentarios.append(random.choice(frases_verdad))
    
    # === ESTRUCTURA Y ORGANIZACIÓN ===
    if ":" in contenido or contenido.count("\n") > 5:
        frases_organizacion = [
            "Tu participacion esta bien organizada y estructurada.",
            "La manera en que organizaste tu informacion es clara.",
            "Aprecio la estructura ordenada de tu aportacion.",
            "Tu trabajo muestra una buena organizacion de ideas."
        ]
        if random.random() < 0.6:  # 60% de probabilidad
            comentarios.append(random.choice(frases_organizacion))
    
    # === PROFUNDIDAD DEL CONTENIDO ===
    longitud_contenido = len(contenido)
    if longitud_contenido > 1800:
        frases_profundidad = [
            "Tu analisis es profundo y completo.",
            "La extension y detalle de tu participacion es destacable.",
            "Tu desarrollo del tema es exhaustivo y bien estructurado.",
            "La profundidad de tu aporte refleja un excelente estudio.",
            "Tu trabajo demuestra una investigacion seria del tema.",
            "El nivel de detalle en tu participacion es admirable."
        ]
        if random.random() < 0.7:  # 70% de probabilidad
            comentarios.append(random.choice(frases_profundidad))
    elif longitud_contenido > 1200:
        frases_buen_desarrollo = [
            "Tu desarrollo del tema es completo.",
            "Tu aportacion tiene un buen nivel de detalle.",
            "El contenido que compartiste es sustancial."
        ]
        if random.random() < 0.5:  # 50% de probabilidad
            comentarios.append(random.choice(frases_buen_desarrollo))
    
    # === MENSAJES FINALES MOTIVACIONALES (más variados) ===
    mensajes_finales = [
        "Tu comprension del tema demuestra un excelente trabajo de estudio. Sigue asi.",
        "Tu participacion refleja dedicacion y esfuerzo. Muy bien.",
        "Excelente trabajo, tu aportacion enriquece el foro. Felicidades.",
        "Tu analisis es muy completo y bien fundamentado. Continua con ese nivel.",
        "Demuestras dominio del tema. Excelente aportacion.",
        "Sigue participando con este nivel de calidad. Felicidades.",
        "Tu esfuerzo es evidente y muy valorado. Excelente.",
        "Continua trabajando con esta dedicacion. Muy bien hecho.",
        "Tu aporte es significativo para el aprendizaje colectivo. Gracias.",
        "Excelente nivel de analisis. Te felicito.",
        "Tu compromiso con el tema es admirable. Adelante.",
        "Muy buen trabajo. Sigue asi.",
        "Tu participacion es de calidad. Felicidades.",
        "Gracias por tu valiosa contribucion.",
        "Tu trabajo refleja profesionalismo. Excelente.",
        "Felicidades por tu dedicacion al tema.",
        "Sigue con ese entusiasmo por aprender.",
        "Tu aportacion es muy valiosa para todos."
    ]
    
    # Agregar mensaje final (75% de probabilidad para no ser tan predecible)
    if random.random() < 0.75:
        comentarios.append(random.choice(mensajes_finales))
    
    # Unir todos los comentarios con espacio
    retroalimentacion = " ".join(comentarios)
    
    # Limpiar el texto para que sea compatible con Moodle
    retroalimentacion = limpiar_texto_para_moodle(retroalimentacion)
    
    return retroalimentacion

def buscar_columna_flexible(df, nombres_posibles):
    """
    Busca una columna de manera flexible, considerando diferentes variaciones de mayúsculas/minúsculas
    y espacios
    """
    columnas_df = df.columns.tolist()
    
    for nombre_buscado in nombres_posibles:
        # Búsqueda exacta
        if nombre_buscado in columnas_df:
            return nombre_buscado
        
        # Búsqueda insensible a mayúsculas/minúsculas
        for col in columnas_df:
            if col.lower() == nombre_buscado.lower():
                return col
        
        # Búsqueda con normalización de espacios
        nombre_normalizado = nombre_buscado.lower().strip()
        for col in columnas_df:
            col_normalizada = col.lower().strip()
            if col_normalizada == nombre_normalizado:
                return col
    
    return None

def mostrar_r4md():
    st.title("🧠 R4MD - Proposiciones Lógicas")
    
    # Tabs para organizar mejor el contenido
    tab_principal = st.tabs(["📤 Calificador Automático de Foros", "💬 Mensajes Simples Excel"])
    
    # ==================== TAB 1: CALIFICADOR AUTOMÁTICO ====================
    with tab_principal[0]:
        st.markdown("### Sistema de Calificación Automática de Foros")
        st.markdown("---")
        
        # Sidebar para configuración (dentro de la tab)
        with st.sidebar:
            st.header("⚙️ Configuración")
            
            # Opción para limpiar historial
            if st.button("🗑️ Limpiar Historial", help="Elimina el registro de alumnos ya calificados"):
                if os.path.exists(HISTORIAL_FILE_R4):
                    os.remove(HISTORIAL_FILE_R4)
                    st.success("Historial eliminado")
                    st.rerun()
            
            # Mostrar estadísticas del historial
            historial = cargar_historial_r4()
            st.metric("Alumnos en historial", len(historial))
            
            st.markdown("---")
            st.markdown("### 📋 Requisitos del foro")
            st.markdown("""
            - ✅ Presentación personal
            - ✅ Definición de proposiciones simples
            - ✅ 5 ejemplos de proposiciones simples
            - ✅ Definición de proposiciones compuestas
            - ✅ 5 ejemplos de proposiciones compuestas
            """)
            
            st.markdown("---")
            st.markdown("### 🎯 Mejoras de Precisión")
            st.info("""
            ✨ **Nueva versión mejorada:**
            - Validación de nombre + apellido
            - Algoritmo de similitud de nombres
            - Reducción de falsos positivos
            - Umbral de coincidencia: 75%
            """)
        
        # Tabs secundarias
        tab1, tab2, tab3, tab4 = st.tabs(["📤 Cargar Archivos", "📊 Resultados", "📜 Historial", "🔍 Debug"])
        
        with tab1:
            st.header("Cargar Archivos para Calificación")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("1️⃣ Archivo Excel")
                excel_file = st.file_uploader(
                    "Sube el archivo de calificaciones (.xlsx)",
                    type=['xlsx'],
                    help="Archivo Excel con la lista de alumnos y sus calificaciones",
                    key="excel_uploader_r4"
                )
                
                if excel_file:
                    st.success("✅ Archivo Excel cargado")
                    try:
                        df = pd.read_excel(excel_file)
                        st.session_state['df_excel_r4'] = df
                        st.info(f"📊 Total de alumnos: {len(df)}")
                        
                        # Verificar que exista la columna necesaria
                        if 'Tarea:R4. Proposiciones lógicas (Real)' in df.columns:
                            st.success("✅ Columna de calificaciones encontrada")
                        else:
                            st.error("❌ No se encontró la columna 'Tarea:R4. Proposiciones lógicas (Real)'")
                        
                        # Verificar columnas de nombre y apellido
                        if 'Nombre' in df.columns and 'Apellido(s)' in df.columns:
                            st.success("✅ Columnas de Nombre y Apellido encontradas")
                        else:
                            st.warning("⚠️ Verifica que existan las columnas 'Nombre' y 'Apellido(s)'")
                        
                        # Mostrar preview
                        with st.expander("Ver preview del Excel"):
                            cols_preview = ['Nombre', 'Apellido(s)']
                            if 'Tarea:R4. Proposiciones lógicas (Real)' in df.columns:
                                cols_preview.append('Tarea:R4. Proposiciones lógicas (Real)')
                            st.dataframe(df[cols_preview].head(10))
                    except Exception as e:
                        st.error(f"Error al leer el Excel: {e}")
            
            with col2:
                st.subheader("2️⃣ Archivo HTML")
                html_file = st.file_uploader(
                    "Sube el HTML del foro (.html)",
                    type=['html'],
                    help="Archivo HTML con las participaciones del foro",
                    key="html_uploader_r4"
                )
                
                if html_file:
                    st.success("✅ Archivo HTML cargado")
                    try:
                        html_content = html_file.read().decode('utf-8')
                        participaciones = extraer_participaciones_html(html_content)
                        
                        # Guardar en session_state
                        st.session_state['participaciones_r4'] = participaciones
                        st.session_state['html_cargado_r4'] = True
                        
                        st.info(f"💬 Participaciones encontradas: {len(participaciones)}")
                        
                        # Mostrar preview
                        with st.expander("Ver preview de participaciones"):
                            if len(participaciones) > 0:
                                for i, p in enumerate(participaciones[:5]):
                                    st.markdown(f"**{p['nombre_completo']}**")
                                    st.caption(f"Primer nombre: {p['primer_nombre']} | Segundo nombre: {p['segundo_nombre']} | Apellidos: {p['apellidos']}")
                                    st.caption(f"Contenido: {p['contenido'][:150]}...")
                                    if i < min(len(participaciones) - 1, 4):
                                        st.markdown("---")
                            else:
                                st.warning("No se encontraron participaciones en el HTML")
                    except Exception as e:
                        st.error(f"Error al leer el HTML: {e}")
                        import traceback
                        st.code(traceback.format_exc())
            
            # Botón de procesamiento
            st.markdown("---")
            
            # Verificar que ambos archivos estén cargados
            archivos_listos = ('df_excel_r4' in st.session_state and 
                               'participaciones_r4' in st.session_state and 
                               st.session_state.get('html_cargado_r4', False))
            
            if archivos_listos:
                if st.button("🚀 Procesar y Generar Retroalimentaciones", type="primary", use_container_width=True):
                    with st.spinner("Procesando participaciones..."):
                        try:
                            # Obtener datos de session_state
                            df = st.session_state['df_excel_r4']
                            participaciones = st.session_state['participaciones_r4']
                            
                            if len(participaciones) == 0:
                                st.error("❌ No se encontraron participaciones para procesar.")
                                st.stop()
                            
                            # Cargar historial
                            historial = cargar_historial_r4()
                            
                            # Crear DataFrame de resultados
                            resultados = []
                            nuevos_calificados = 0
                            ya_calificados = 0
                            no_encontrados = 0
                            
                            # Para debug: guardar información de matching
                            debug_info = []
                            
                            # Procesar cada participación
                            progress_bar = st.progress(0)
                            status_text = st.empty()
                            
                            for i, p in enumerate(participaciones):
                                progress_bar.progress((i + 1) / len(participaciones))
                                status_text.text(f"Procesando: {p['nombre_completo']}")
                                
                                # Buscar alumno en Excel con validación mejorada
                                idx = buscar_alumno_en_excel(
                                    df, 
                                    p['nombre_completo'],
                                    p['primer_nombre'], 
                                    p['segundo_nombre'],
                                    p['apellidos']
                                )
                                
                                if idx is not None:
                                    nombre_excel = str(df.loc[idx, 'Nombre'])
                                    apellido_excel = str(df.loc[idx, 'Apellido(s)'])
                                    nombre_completo_excel = f"{nombre_excel} {apellido_excel}"
                                    primer_nombre_excel = nombre_excel.split()[0]  # Extraer primer nombre
                                    calificacion_actual = df.loc[idx, 'Tarea:R4. Proposiciones lógicas (Real)']
                                    
                                    # Guardar info de debug
                                    similitud = similitud_nombres(p['nombre_completo'], nombre_completo_excel)
                                    debug_info.append({
                                        'HTML': p['nombre_completo'],
                                        'Excel': nombre_completo_excel,
                                        'Similitud': f"{similitud:.2%}",
                                        'Match': '✅'
                                    })
                                    
                                    # Verificar si tiene "-" (sin calificar)
                                    necesita_calificacion = (pd.isna(calificacion_actual) or 
                                                            str(calificacion_actual).strip() == "-" or 
                                                            str(calificacion_actual).strip() == "")
                                    
                                    # Verificar si ya está en historial
                                    en_historial = nombre_completo_excel in historial
                                    
                                    if necesita_calificacion and not en_historial:
                                        # Generar retroalimentación con nombre completo y primer nombre
                                        retroalimentacion = generar_retroalimentacion_r4(nombre_completo_excel, primer_nombre_excel, p['contenido'])
                                        
                                        # Agregar al historial
                                        historial[nombre_completo_excel] = {
                                            'fecha': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                            'retroalimentacion': retroalimentacion
                                        }
                                        
                                        resultados.append({
                                            'Nombre': nombre_completo_excel,
                                            'Retroalimentación': retroalimentacion
                                        })
                                        
                                        nuevos_calificados += 1
                                    else:
                                        ya_calificados += 1
                                else:
                                    no_encontrados += 1
                                    debug_info.append({
                                        'HTML': p['nombre_completo'],
                                        'Excel': 'No encontrado',
                                        'Similitud': 'N/A',
                                        'Match': '❌'
                                    })
                            
                            progress_bar.empty()
                            status_text.empty()
                            
                            # Guardar historial actualizado
                            guardar_historial_r4(historial)
                            
                            # Crear DataFrame de resultados
                            if resultados:
                                df_resultados = pd.DataFrame(resultados)
                            else:
                                df_resultados = pd.DataFrame(columns=['Nombre', 'Retroalimentación'])
                            
                            # Guardar en session_state
                            st.session_state['df_resultados_r4'] = df_resultados
                            st.session_state['nuevos_calificados_r4'] = nuevos_calificados
                            st.session_state['ya_calificados_r4'] = ya_calificados
                            st.session_state['no_encontrados_r4'] = no_encontrados
                            st.session_state['total_participaciones_r4'] = len(participaciones)
                            st.session_state['debug_info_r4'] = debug_info
                            
                            st.success(f"✅ Procesamiento completado!")
                            st.balloons()
                            
                        except Exception as e:
                            st.error(f"Error durante el procesamiento: {e}")
                            import traceback
                            st.code(traceback.format_exc())
            else:
                st.info("👆 Por favor, carga ambos archivos (Excel y HTML) para continuar")
        
        with tab2:
            st.header("Resultados del Procesamiento")
            
            if 'df_resultados_r4' in st.session_state:
                # Mostrar métricas
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("🆕 Nuevos Calificados", st.session_state['nuevos_calificados_r4'])
                with col2:
                    st.metric("✅ Ya Calificados", st.session_state['ya_calificados_r4'])
                with col3:
                    st.metric("❌ No Encontrados", st.session_state['no_encontrados_r4'])
                with col4:
                    st.metric("💬 Total Participaciones", st.session_state['total_participaciones_r4'])
                
                st.markdown("---")
                
                # Mostrar tabla de resultados
                df_resultado = st.session_state['df_resultados_r4']
                
                if len(df_resultado) > 0:
                    st.subheader(f"📋 Retroalimentaciones Generadas ({len(df_resultado)})")
                    st.dataframe(
                        df_resultado,
                        use_container_width=True,
                        hide_index=True
                    )
                    
                    # Mostrar retroalimentaciones individuales
                    st.markdown("---")
                    st.subheader("💬 Detalle de Retroalimentaciones")
                    
                    for idx, row in df_resultado.iterrows():
                        with st.expander(f"👤 {row['Nombre']}"):
                            st.info(row['Retroalimentación'])
                    
                    # Botón de descarga
                    st.markdown("---")
                    st.subheader("💾 Descargar Resultados")
                    
                    # Preparar archivo para descarga
                    output = io.BytesIO()
                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                        df_resultado.to_excel(writer, index=False, sheet_name='Retroalimentaciones')
                    output.seek(0)
                    
                    st.download_button(
                        label="📥 Descargar Excel con Retroalimentaciones",
                        data=output,
                        file_name=f"Retroalimentaciones_R4_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        type="primary",
                        use_container_width=True
                    )
                    
                    # Sección de texto plano para Moodle
                    st.markdown("---")
                    st.subheader("📋 Texto Plano para Moodle")
                    st.info("""
                    💡 **Cómo usar:** Selecciona el texto del cuadro de abajo, cópialo (Ctrl+C o Cmd+C) 
                    y pégalo directamente en Moodle. Este formato está limpio y no causará errores de JSON.
                    """)
                    
                    # Crear formato de texto plano
                    texto_plano = ""
                    for idx, row in df_resultado.iterrows():
                        nombre = row['Nombre']
                        retro = row['Retroalimentación']
                        
                        texto_plano += f"{nombre}\n"
                        texto_plano += f"{retro}\n"
                        texto_plano += "\n" + "="*80 + "\n\n"
                    
                    # Mostrar en un text_area para fácil copiado
                    st.text_area(
                        "Selecciona todo el texto (Ctrl+A) y copia (Ctrl+C):",
                        texto_plano,
                        height=400,
                        key="texto_plano_moodle_r4"
                    )
                    
                    # Botón para descargar como TXT
                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button(
                            label="📥 Descargar como TXT",
                            data=texto_plano.encode('utf-8'),
                            file_name=f"Retroalimentaciones_R4_Texto_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain",
                            use_container_width=True
                        )
                    with col2:
                        # Botón para copiar al portapapeles (información)
                        st.info("💡 Tip: Usa Ctrl+A para seleccionar todo el texto arriba")
                    
                    # Opción alternativa: formato simple sin separadores
                    st.markdown("---")
                    st.subheader("📝 Formato Simple (Alternativo)")
                    st.caption("Si el formato anterior causa problemas, usa este formato más simple:")
                    
                    texto_simple = ""
                    for idx, row in df_resultado.iterrows():
                        texto_simple += f"{row['Nombre']}: {row['Retroalimentación']}\n\n"
                    
                    st.text_area(
                        "Texto simple sin formato:",
                        texto_simple,
                        height=300,
                        key="texto_simple_moodle_r4"
                    )
                    
                    st.download_button(
                        label="📥 Descargar formato simple como TXT",
                        data=texto_simple.encode('utf-8'),
                        file_name=f"Retroalimentaciones_R4_Simple_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                else:
                    st.info("📝 No hay retroalimentaciones nuevas generadas en este procesamiento.")
                    st.markdown("""
                    **Posibles razones:**
                    - Todos los alumnos ya tienen calificación (no tienen "-")
                    - Todos ya fueron calificados anteriormente (están en historial)
                    - No se encontraron coincidencias entre el foro y el Excel
                    """)
            else:
                st.info("👆 Carga los archivos en la pestaña 'Cargar Archivos' y procesa para ver resultados.")
        
        with tab3:
            st.header("Historial de Calificaciones")
            
            historial = cargar_historial_r4()
            
            if historial:
                st.success(f"📊 Total de alumnos en historial: {len(historial)}")
                
                # Convertir historial a DataFrame
                historial_data = []
                for nombre, datos in historial.items():
                    historial_data.append({
                        'Nombre': nombre,
                        'Fecha Calificación': datos['fecha'],
                        'Retroalimentación': datos['retroalimentacion']
                    })
                
                df_historial = pd.DataFrame(historial_data)
                
                # Buscar en historial
                buscar = st.text_input("🔍 Buscar alumno en historial:", placeholder="Escribe el nombre...")
                
                if buscar:
                    df_filtrado = df_historial[df_historial['Nombre'].str.contains(buscar, case=False, na=False)]
                    st.dataframe(df_filtrado, use_container_width=True, hide_index=True)
                    
                    # Mostrar retroalimentación completa
                    if len(df_filtrado) > 0:
                        for idx, row in df_filtrado.iterrows():
                            nombre = row['Nombre']
                            with st.expander(f"Ver retroalimentación completa de {nombre}"):
                                st.info(historial[nombre]['retroalimentacion'])
                                st.caption(f"📅 Fecha: {historial[nombre]['fecha']}")
                else:
                    st.dataframe(df_historial, use_container_width=True, hide_index=True)
                
                # Exportar historial
                st.markdown("---")
                historial_json = json.dumps(historial, ensure_ascii=False, indent=2)
                st.download_button(
                    label="📥 Descargar Historial (JSON)",
                    data=historial_json,
                    file_name=f"historial_calificaciones_{datetime.now().strftime('%Y%m%d')}.json",
                    mime="application/json"
                )
                
            else:
                st.info("📝 No hay registros en el historial aún. Procesa algunas participaciones para comenzar.")
        
        with tab4:
            st.header("🔍 Información de Debug - Matching de Nombres")
            
            if 'debug_info_r4' in st.session_state and len(st.session_state['debug_info_r4']) > 0:
                st.info("""
                Esta pestaña muestra la información detallada del proceso de matching entre los nombres del HTML 
                y los nombres del Excel. Úsala para verificar que los matches sean correctos y diagnosticar falsos positivos.
                """)
                
                df_debug = pd.DataFrame(st.session_state['debug_info_r4'])
                
                # Filtros
                col1, col2 = st.columns(2)
                with col1:
                    filtro_match = st.selectbox(
                        "Filtrar por resultado:",
                        ["Todos", "Solo matches (✅)", "Solo no encontrados (❌)"]
                    )
                
                with col2:
                    filtro_similitud = st.slider(
                        "Similitud mínima a mostrar:",
                        0.0, 1.0, 0.0, 0.05,
                        help="Solo aplica a los matches encontrados"
                    )
                
                # Aplicar filtros
                df_filtrado = df_debug.copy()
                
                if filtro_match == "Solo matches (✅)":
                    df_filtrado = df_filtrado[df_filtrado['Match'] == '✅']
                elif filtro_match == "Solo no encontrados (❌)":
                    df_filtrado = df_filtrado[df_filtrado['Match'] == '❌']
                
                # Filtrar por similitud (solo para matches)
                if filtro_similitud > 0:
                    df_filtrado = df_filtrado[
                        (df_filtrado['Match'] == '❌') | 
                        (df_filtrado['Similitud'].apply(lambda x: float(x.strip('%')) / 100 if x != 'N/A' else 0) >= filtro_similitud)
                    ]
                
                st.dataframe(
                    df_filtrado,
                    use_container_width=True,
                    hide_index=True
                )
                
                # Estadísticas
                st.markdown("---")
                st.subheader("📈 Estadísticas de Matching")
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    total_matches = len(df_debug[df_debug['Match'] == '✅'])
                    st.metric("Total de Matches", total_matches)
                
                with col2:
                    total_no_encontrados = len(df_debug[df_debug['Match'] == '❌'])
                    st.metric("No Encontrados", total_no_encontrados)
                
                with col3:
                    if total_matches > 0:
                        similitudes = [float(x.strip('%')) / 100 for x in df_debug[df_debug['Match'] == '✅']['Similitud']]
                        similitud_promedio = sum(similitudes) / len(similitudes)
                        st.metric("Similitud Promedio", f"{similitud_promedio:.1%}")
                
                # Lista de no encontrados
                if total_no_encontrados > 0:
                    st.markdown("---")
                    st.subheader("⚠️ Nombres del HTML no encontrados en Excel")
                    st.warning("""
                    Estos participantes del foro no se encontraron en el Excel. Posibles causas:
                    - El nombre en el foro es diferente al del Excel
                    - El alumno no está registrado en el Excel
                    - Hay errores de ortografía en alguno de los dos archivos
                    """)
                    no_encontrados_lista = df_debug[df_debug['Match'] == '❌']['HTML'].tolist()
                    for nombre in no_encontrados_lista:
                        st.text(f"• {nombre}")
                
            else:
                st.info("👆 Procesa archivos en la pestaña 'Cargar Archivos' para ver información de debug.")
    
    # ==================== TAB 2: MENSAJES SIMPLES (ORIGINAL) ====================
    with tab_principal[1]:
        st.markdown("### Mensajes Simples desde Excel")
        st.markdown("---")
        
        mensajes_r4 = [
            "Buen día {nombre}. He tenido la oportunidad de revisar tu participación en el foro y quiero felicitarte, ya que has abordado todos los puntos de manera adecuada, cumpliendo con los criterios de la rúbrica. Ahora, aguardamos los comentarios de tus compañeros para enriquecer el intercambio. Te sugiero considerar sus observaciones y sacar provecho de esta oportunidad. ¡Saludos!",
            
            "Hola {nombre}, qué gusto saludarte. Revisé tu trabajo en el foro y quiero felicitarte por cumplir con los puntos solicitados en la rúbrica. Ahora esperemos la retroalimentación de tus compañeros, ya que el foro está diseñado para promover este intercambio de ideas. Aprovecha los comentarios recibidos para potenciar tu aprendizaje. Saludos.",
            
            "Gracias por tu aporte {nombre}. He revisado con detalle tu participación en el foro y quiero reconocerte el haber cumplido con todos los criterios establecidos. Ahora, esperamos las observaciones de tus compañeros, que enriquecerán la discusión y te brindarán nuevos puntos de vista. Aprovecha esta oportunidad para fortalecer tus conocimientos. Saludos cordiales.",
            
            "Excelente trabajo {nombre}. Al revisar tu contribución en el foro, pude ver que has cumplido con todos los aspectos solicitados en la rúbrica, ¡felicidades! Ahora queda por esperar los comentarios de tus compañeros, quienes podrán ofrecerte nuevas perspectivas. Considera sus observaciones para sacar el mayor provecho de esta actividad. Saludos.",
            
            "¿Qué tal? {nombre}. Muy bien hecho. Tu participación en el foro ha sido revisada, y es evidente que has cumplido con los puntos solicitados de forma satisfactoria. Ahora, espera la retroalimentación de tus compañeros, ya que el intercambio de ideas es el objetivo de este espacio. Aprovecha sus comentarios para fortalecer tu aprendizaje. ¡Saludos!"
        ]
        
        excel_file = st.file_uploader("📊 Carga el archivo Excel", type=["xlsx"], key="excel_r4_simple")
        
        if excel_file:
            try:
                df = pd.read_excel(excel_file)
                
                # Mostrar información del archivo
                st.info(f"📋 Archivo cargado: {len(df)} filas, {len(df.columns)} columnas")
                
                # Mostrar columnas disponibles
                with st.expander("👁️ Ver columnas disponibles"):
                    st.write(list(df.columns))
                
                # Buscar columnas de manera flexible
                nombres_columna_objetivo = [
                    "Tarea:R4. Proposiciones lógicas (Real)",
                    "Tarea: R4. Proposiciones lógicas (Real)",
                    "Tarea:R4.Proposiciones lógicas (Real)"
                ]
                
                nombres_columna_nombre = [
                    "Nombre",
                    "nombre", 
                    "NOMBRE",
                    "Nombre completo",
                    "nombre completo"
                ]
                
                columna_objetivo = buscar_columna_flexible(df, nombres_columna_objetivo)
                columna_nombre = buscar_columna_flexible(df, nombres_columna_nombre)
                
                if columna_objetivo:
                    st.success(f"✅ Columna objetivo encontrada: '{columna_objetivo}'")
                    
                    # Filtrar filas con "-"
                    filas_con_guion = df[df[columna_objetivo] == "-"]
                    
                    if len(filas_con_guion) > 0:
                        st.info(f"🔍 Encontradas {len(filas_con_guion)} filas con '-'")
                        
                        if columna_nombre:
                            st.success(f"✅ Columna nombre encontrada: '{columna_nombre}'")
                            
                            # Obtener nombres
                            nombres = filas_con_guion[columna_nombre].tolist()
                            
                            # Limpiar nombres (quitar espacios extra, NaN, etc.)
                            nombres_limpios = []
                            for nombre in nombres:
                                if pd.notna(nombre) and str(nombre).strip():
                                    nombres_limpios.append(str(nombre).strip())
                            
                            if nombres_limpios:
                                # Crear mensajes balanceados
                                mensajes_finales = []
                                datos_para_excel = []
                                
                                st.markdown("---")
                                st.subheader("📝 Mensajes Generados")
                                
                                for i, nombre in enumerate(nombres_limpios):
                                    # Usar módulo para distribuir mensajes de manera equilibrada
                                    mensaje_idx = i % len(mensajes_r4)
                                    mensaje_completo = mensajes_r4[mensaje_idx].format(nombre=nombre)
                                    mensajes_finales.append(mensaje_completo)
                                    
                                    # Datos para Excel (nombre y mensaje en columnas separadas)
                                    datos_para_excel.append({
                                        'Nombre': nombre,
                                        'Mensaje': mensaje_completo
                                    })
                                    
                                    # Mostrar cada mensaje con su botón individual
                                    with st.container():
                                        st.markdown(f"**{i+1}. {nombre}**")
                                        
                                        # Mostrar el mensaje en un área de texto pequeña
                                        st.text_area(
                                            f"Mensaje para {nombre}:", 
                                            value=mensaje_completo, 
                                            height=120, 
                                            key=f"mensaje_simple_{i}",
                                            label_visibility="collapsed"
                                        )
                                        
                                        # Botón para copiar mensaje individual
                                        if st.button(f"📋 Copiar mensaje de {nombre}", key=f"copy_individual_simple_{i}"):
                                            components.html(copy_to_clipboard_js(mensaje_completo), height=0)
                                            st.success(f"✅ ¡Mensaje de {nombre} copiado!")
                                        
                                        st.markdown("---")
                                
                                # Crear DataFrame para Excel con estructura solicitada
                                df_resultado = pd.DataFrame(datos_para_excel)
                                
                                st.success(f"✅ Procesados {len(mensajes_finales)} mensajes")
                                
                                # Mostrar DataFrame resultado
                                st.subheader("📊 Vista previa del Excel")
                                st.dataframe(df_resultado)
                                
                                # Botones principales
                                col1, col2, col3 = st.columns(3)
                                
                                with col1:
                                    # Copiar todos los mensajes (solo el contenido, sin nombres)
                                    texto_todos_mensajes = "\n\n".join(mensajes_finales)
                                    if st.button("📋 Copiar TODOS los mensajes", type="primary", key="copy_all_simple"):
                                        components.html(copy_to_clipboard_js(texto_todos_mensajes), height=0)
                                        st.success("✅ ¡Todos los mensajes copiados!")
                                
                                with col2:
                                    # Descargar Excel con estructura nombre|mensaje
                                    output = io.BytesIO()
                                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                                        df_resultado.to_excel(writer, index=False, sheet_name='Mensajes_R4')
                                    
                                    st.download_button(
                                        "📥 Descargar Excel",
                                        data=output.getvalue(),
                                        file_name="mensajes_r4.xlsx",
                                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                    )
                                
                                with col3:
                                    # Descargar solo mensajes como TXT
                                    st.download_button(
                                        "📄 Descargar mensajes TXT",
                                        data=texto_todos_mensajes,
                                        file_name="mensajes_r4.txt",
                                        mime="text/plain"
                                    )
                                
                                # Mostrar distribución de mensajes
                                with st.expander("📊 Distribución de mensajes"):
                                    distribucion = {}
                                    for i in range(len(mensajes_finales)):
                                        mensaje_tipo = f"Mensaje {(i % len(mensajes_r4)) + 1}"
                                        distribucion[mensaje_tipo] = distribucion.get(mensaje_tipo, 0) + 1
                                    
                                    for tipo, cantidad in distribucion.items():
                                        st.write(f"{tipo}: {cantidad} veces")
                            
                            else:
                                st.warning("⚠️ No se encontraron nombres válidos en las filas con '-'")
                        
                        else:
                            st.error(f"❌ No se encontró ninguna columna de nombres")
                            st.write("**Columnas buscadas:** ", nombres_columna_nombre)
                            st.write("**Columnas disponibles:** ", list(df.columns))
                            
                            # Sugerir columnas similares
                            st.write("**💡 Sugerencias de columnas que podrían contener nombres:**")
                            for col in df.columns:
                                if any(palabra in col.lower() for palabra in ['nombre', 'name', 'alumno', 'estudiante']):
                                    st.write(f"   - {col}")
                    
                    else:
                        st.warning("⚠️ No se encontraron filas con '-' en la columna objetivo")
                        
                        # Mostrar valores únicos de la columna objetivo para debug
                        with st.expander("🔍 Ver valores únicos en la columna objetivo"):
                            valores_unicos = df[columna_objetivo].value_counts()
                            st.write(valores_unicos)
                
                else:
                    st.error(f"❌ No se encontró la columna objetivo")
                    st.write("**Columnas buscadas:** ", nombres_columna_objetivo)
                    st.write("**Columnas disponibles:** ", list(df.columns))
                    
                    # Sugerir columnas similares
                    st.write("**💡 Sugerencias de columnas que podrían ser la objetivo:**")
                    for col in df.columns:
                        if any(palabra in col.lower() for palabra in ['tarea', 'r4', 'proposiciones', 'logicas']):
                            st.write(f"   - {col}")
            
            except Exception as e:
                st.error(f"❌ Error al procesar el archivo Excel: {str(e)}")

# ==================== R7MD - MENSAJES PREDEFINIDOS ====================

def mostrar_r7md():
    st.title("💬 R7MD - Mensajes Predefinidos")
    
    # Inicializar session_state PRIMERO antes de cualquier otra cosa
    if 'mensaje_correcto_aleatorio_r7' not in st.session_state:
        st.session_state['mensaje_correcto_aleatorio_r7'] = random.randint(0, 4)
    
    if 'nombre_anterior_r7' not in st.session_state:
        st.session_state['nombre_anterior_r7'] = ''
    
    # Campo para capturar el nombre del alumno
    st.markdown("---")
    st.subheader("👤 Personalización de Mensajes")
    nombre_alumno = st.text_input(
        "Ingresa el nombre del alumno para personalizar los mensajes:",
        placeholder="Ejemplo: Juan, María, Carlos...",
        key="nombre_alumno_r7"
    )
    
    # Obtener el nombre limpio (evitar problemas con espacios)
    nombre_limpio = nombre_alumno.strip() if nombre_alumno else ""
    
    # Si cambió el nombre, generar nuevo mensaje aleatorio
    if nombre_limpio != st.session_state['nombre_anterior_r7']:
        st.session_state['mensaje_correcto_aleatorio_r7'] = random.randint(0, 4)
        st.session_state['nombre_anterior_r7'] = nombre_limpio
    
    # Mostrar estado de personalización
    if nombre_limpio:
        st.success(f"✅ Mensajes personalizados para: **{nombre_limpio}**")
    else:
        st.info("ℹ️ Los mensajes se mostrarán con el placeholder {{nombre}}. Ingresa un nombre para personalizarlos.")
    
    st.markdown("---")
    
    # Mensajes CORRECTOS (5 variantes)
    mensajes_correcto = [
        """Excelente trabajo {nombre}, se requiere poner en práctica todo el conocimiento del curso para lograr resolver el ejercicio como lo has hecho, identificas de manera adecuada todos los elementos solicitados, continúa así.

Ha sido un gusto acompañarte en este proceso de aprendizaje. Te deseo mucho éxito en tus subsecuentes módulos.

Saludos.""",

        """Felicidades {nombre}, tu desempeño en este ejercicio ha sido sobresaliente. Se nota que has comprendido a profundidad todos los conceptos trabajados durante el curso y los has aplicado correctamente. Sigue con ese nivel de dedicación.

Ha sido un placer acompañarte en este proceso de aprendizaje. Te deseo mucho éxito en tus subsecuentes módulos.

Saludos.""",

        """Muy bien hecho {nombre}, tu trabajo demuestra un dominio completo de los temas abordados. La forma en que has resuelto el ejercicio refleja que has aprovechado al máximo el curso. Te felicito por tu esfuerzo y compromiso.

Ha sido un honor acompañarte en este proceso de aprendizaje. Te deseo mucho éxito en tus subsecuentes módulos.

Saludos.""",

        """Extraordinario trabajo {nombre}, has logrado integrar todos los conocimientos del curso de manera impecable. Tu solución es precisa y muestra claramente que has alcanzado los objetivos de aprendizaje. Continúa con ese excelente desempeño.

Ha sido un privilegio acompañarte en este proceso de aprendizaje. Te deseo mucho éxito en tus subsecuentes módulos.

Saludos.""",

        """Perfecto {nombre}, tu ejercicio refleja una comprensión integral de todos los temas del curso. Has demostrado habilidad para aplicar los conceptos de manera correcta y efectiva. Felicidades por tu logro.

Ha sido gratificante acompañarte en este proceso de aprendizaje. Te deseo mucho éxito en tus subsecuentes módulos.

Saludos."""
    ]
    
    # Mensajes INCORRECTOS (4 mensajes)
    mensajes_incorrecto = [
        """Buen trabajo {nombre}, lo que corresponde a tu primera tabla es correcto, identificas de manera adecuada las propiedades, la segunda tabla no se ha realizado, de ahí tu calificación, te dejo un video que he realizado con el objetivo de poder darte claridad para resolver el ejercicio.

https://youtu.be/naYR2TQ84L0

Corrige y reenvía.

Saludos.""",

        """Buen trabajo {nombre}, lo que corresponde a tu primera tabla es correcto, identificas de manera adecuada las propiedades, en la segunda tabla la que corresponde al diagrama de Hasse, es correcta hasta el paso 3, ya que en el paso 4, a pesar que identificas de manera correcta cada una de las relaciones transitivas, hay un cambio de dirección de la arista de "c" a "b", ya que la dirección en un paso anterior lo manejas de "b" a "c", de ahí la calificación, si pudieras argumentar dicho cambio de dirección podría corregir la calificación, quedo al pendiente.

Aquí un video que te puede ayudar.

https://youtu.be/WTGkSBsLX34

Saludos.""",

        """Buen trabajo {nombre}, lo que corresponde a tu primera tabla es correcto, identificas de manera adecuada las propiedades, en la segunda tabla la que corresponde al diagrama de Hasse, es correcta hasta el paso 2, ya que en el paso 3, no identificas en su totalidad las relaciones transitivas, situación que te lleva al error en tu diagrama final, te dejo un video que he realizado con el objetivo de poder darte claridad para resolver el ejercicio.

https://youtu.be/naYR2TQ84L0

Corrige y reenvía.

Saludos.""",

        """Buen trabajo {nombre}, la primera tabla es correcta, en la parte que corresponde al dígrafo faltó eliminar la totalidad de las relaciones transitivas, hecho que no te permite alcanzar el 100% de la calificación.

Te dejo la resolución del ejercicio y quedo a disposición por si hubiera alguna duda más, aprovecho para preguntar, con todo respeto ¿Viste el video que te envíe en la realimentación anterior?

Saludos.

https://youtu.be/WTGkSBsLX34"""
    ]
    
    # Mensajes ALTERNOS (3 mensajes)
    mensajes_alternos = [
        """Buen trabajo {nombre}, un detalle en el dígrafo de la primera tabla, en particular en la relación transitiva, ya que no corresponde la notación matemática y el dígrafo, se otorga la mayor calificación esperando tomes en consideración la observación.

Saludos.""",

        """Buen trabajo {nombre}, lo que corresponde a tu primera tabla es correcto, identificas de manera adecuada las propiedades, en la segunda tabla la que corresponde al diagrama de Hasse, es correcta hasta el paso 3, ya que en el paso 4, estás realizando un acomodo incorrecto, situación que te lleva al error en tu diagrama final, te dejo un video que he realizado con el objetivo de poder darte claridad para resolver el ejercicio. Esperando tomes en consideración la recomendación, para evitar suspicacia en futuros trabajo, se asigna la mayor calificación.

https://youtu.be/WTGkSBsLX34

Éxito en tus subsecuentes retos.

Ha sido un placer acompañarte en este proceso de aprendizaje. Te deseo mucho éxito en tus subsecuentes módulos.

Saludos.""",

        """Buen trabajo {nombre}, desafortunadamente este trabajo lo he visto en entregas anteriores, de hecho veo que es prácticamente el mismo trabajo que compañeros tuyos están entregando, estoy llegando a la conclusión que es un trabajo que aparece en Internet. Me hubiera gustado que generarás tu propio diseño y con ello, adueñarte del conocimiento, ya que en este caso, solo terminas copiando y pegando, sin reflexionar lo que implica este ejercicio. Por cualquier duda quedo a disposición.

Saludos."""
    ]
    
    # Mostrar en tres columnas
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.subheader("🟢 CORRECTOS")
        st.info("💡 El primer mensaje mostrado es aleatorio para cada alumno")
        
        # Obtener el índice del mensaje "destacado" (el primero que se ve)
        indice_destacado = st.session_state['mensaje_correcto_aleatorio_r7']
        
        # Mostrar primero el mensaje aleatorio seleccionado
        mensaje_destacado = mensajes_correcto[indice_destacado]
        if nombre_limpio:
            mensaje_destacado_personalizado = mensaje_destacado.replace("{nombre}", nombre_limpio)
        else:
            mensaje_destacado_personalizado = mensaje_destacado
        
        st.markdown(f"**⭐ Mensaje Sugerido (Variante {indice_destacado + 1}):**")
        st.text_area(
            "Mensaje para copiar:", 
            value=mensaje_destacado_personalizado, 
            height=250, 
            key=f"mensaje_destacado_correcto_{nombre_limpio}"  # Key única que incluye el nombre
        )
        
        if st.button("📋 Copiar Mensaje Sugerido", key="copy_destacado", type="primary", use_container_width=True):
            components.html(copy_to_clipboard_js(mensaje_destacado_personalizado), height=0)
            st.success("✅ Mensaje copiado!")
        
        # Mostrar los demás mensajes en expandibles
        st.markdown("---")
        st.markdown("**Otras variantes disponibles:**")
        for i, mensaje in enumerate(mensajes_correcto, 1):
            # Saltar el mensaje destacado
            if (i - 1) == indice_destacado:
                continue
                
            # Reemplazar {nombre} con el nombre capturado si existe
            if nombre_limpio:
                mensaje_personalizado = mensaje.replace("{nombre}", nombre_limpio)
            else:
                mensaje_personalizado = mensaje
            
            with st.expander(f"Mensaje {i} - Correcto"):
                st.text_area(
                    f"Mensaje {i}", 
                    value=mensaje_personalizado, 
                    height=200, 
                    key=f"correcto_{i}_{nombre_limpio}"  # Key única
                )
                if st.button(f"📋 Copiar Mensaje {i}", key=f"copy_correcto_{i}"):
                    components.html(copy_to_clipboard_js(mensaje_personalizado), height=0)
                    st.success(f"✅ Mensaje {i} copiado!")
    
    with col2:
        st.subheader("🔴 INCORRECTOS")
        for i, mensaje in enumerate(mensajes_incorrecto, 1):
            # Reemplazar {nombre} con el nombre capturado si existe
            if nombre_limpio:
                mensaje_personalizado = mensaje.replace("{nombre}", nombre_limpio)
            else:
                mensaje_personalizado = mensaje
            
            with st.expander(f"Mensaje {i} - Incorrecto"):
                st.text_area(
                    f"Mensaje {i}", 
                    value=mensaje_personalizado, 
                    height=250, 
                    key=f"incorrecto_{i}_{nombre_limpio}"  # Key única
                )
                if st.button(f"📋 Copiar Mensaje {i}", key=f"copy_incorrecto_{i}"):
                    components.html(copy_to_clipboard_js(mensaje_personalizado), height=0)
                    st.success(f"✅ Mensaje {i} copiado!")
    
    with col3:
        st.subheader("🟡 ALTERNOS")
        for i, mensaje in enumerate(mensajes_alternos, 1):
            # Reemplazar {nombre} con el nombre capturado si existe
            if nombre_limpio:
                mensaje_personalizado = mensaje.replace("{nombre}", nombre_limpio)
            else:
                mensaje_personalizado = mensaje
            
            with st.expander(f"Mensaje {i} - Alterno"):
                st.text_area(
                    f"Mensaje {i}", 
                    value=mensaje_personalizado, 
                    height=250, 
                    key=f"alterno_{i}_{nombre_limpio}"  # Key única
                )
                if st.button(f"📋 Copiar Mensaje {i}", key=f"copy_alterno_{i}"):
                    components.html(copy_to_clipboard_js(mensaje_personalizado), height=0)
                    st.success(f"✅ Mensaje {i} copiado!")
    
    # Botones para copiar todos los mensajes de cada categoría
    st.markdown("---")
    st.subheader("📋 Acciones Rápidas")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("**🟢 Mensajes Correctos**")
        if st.button("📋 Copiar TODOS CORRECTOS", type="secondary", use_container_width=True):
            # Personalizar todos los mensajes antes de copiar
            if nombre_limpio:
                mensajes_personalizados = [msg.replace("{nombre}", nombre_limpio) for msg in mensajes_correcto]
            else:
                mensajes_personalizados = mensajes_correcto
            todos_correctos = "\n\n" + "="*50 + "\n\n".join([f"MENSAJE {i+1} - CORRECTO:\n\n{msg}" for i, msg in enumerate(mensajes_personalizados)])
            components.html(copy_to_clipboard_js(todos_correctos), height=0)
            st.success("✅ Todos los mensajes correctos copiados!")
    
    with col2:
        st.markdown("**🔴 Mensajes Incorrectos**")
        if st.button("📋 Copiar TODOS INCORRECTOS", type="secondary", use_container_width=True):
            # Personalizar todos los mensajes antes de copiar
            if nombre_limpio:
                mensajes_personalizados = [msg.replace("{nombre}", nombre_limpio) for msg in mensajes_incorrecto]
            else:
                mensajes_personalizados = mensajes_incorrecto
            todos_incorrectos = "\n\n" + "="*50 + "\n\n".join([f"MENSAJE {i+1} - INCORRECTO:\n\n{msg}" for i, msg in enumerate(mensajes_personalizados)])
            components.html(copy_to_clipboard_js(todos_incorrectos), height=0)
            st.success("✅ Todos los mensajes incorrectos copiados!")
    
    with col3:
        st.markdown("**🟡 Mensajes Alternos**")
        if st.button("📋 Copiar TODOS ALTERNOS", type="secondary", use_container_width=True):
            # Personalizar todos los mensajes antes de copiar
            if nombre_limpio:
                mensajes_personalizados = [msg.replace("{nombre}", nombre_limpio) for msg in mensajes_alternos]
            else:
                mensajes_personalizados = mensajes_alternos
            todos_alternos = "\n\n" + "="*50 + "\n\n".join([f"MENSAJE {i+1} - ALTERNO:\n\n{msg}" for i, msg in enumerate(mensajes_personalizados)])
            components.html(copy_to_clipboard_js(todos_alternos), height=0)
            st.success("✅ Todos los mensajes alternos copiados!")

# ==================== NAVEGACIÓN PRINCIPAL ====================

if menu_option == "R3MD - Conjuntos":
    mostrar_r3md()
elif menu_option == "R4MD - Proposiciones Lógicas":
    mostrar_r4md()
elif menu_option == "R7MD - Mensajes Predefinidos":
    mostrar_r7md()

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666;'>
    <p>🎓 Sistema de Retroalimentación | Matemáticas Discretas</p>
    <p style='font-size: 0.8em;'>Versión V4 + Mejoras V5 | R3MD con delimitadores mixtos y filtrado mejorado</p>
</div>
""", unsafe_allow_html=True)
